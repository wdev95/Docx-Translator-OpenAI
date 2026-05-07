from __future__ import annotations

import json
import hashlib
import math
import os
import re
import shutil
import time
import traceback
import threading
import tkinter as tk
from difflib import SequenceMatcher
import darkdetect
import svc_ttk
import zipfile
from copy import deepcopy
from datetime import datetime
from queue import Empty, Queue
from pathlib import Path
from typing import Any, Callable, Iterable
from tkinter import filedialog, ttk
from urllib import error as urllib_error
from urllib import request as urllib_request

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from openai import OpenAI
import sys
if sys.platform == "win32":
	from PyTaskbar import TaskbarProgress, ProgressType
else:
	TaskbarProgress = None  # type: ignore[assignment,misc]
	class ProgressType:  # type: ignore[no-redef]
		NORMAL = 0
		ERROR = 1
		PAUSED = 2
		INDETERMINATE = 3
		NOPROGRESS = 4

TARGET_STYLE_NAMES = {"Standard", "Normal"}
THIRD_PARTY_PRICING_API_ENDPOINTS = (
	"https://api.openaipricing.com/openai/text_tokens",
	"https://openrouter.ai/api/v1/models",
)
FX_RATE_ENDPOINTS = (
	"https://open.er-api.com/v6/latest/USD",
	"https://api.frankfurter.app/latest?from=USD&to=EUR",
)
DEFAULT_LANGUAGE_CODES = ["de", "en-US", "en-GB", "it", "es-ES", "es-419", "fr"]
DEFAULT_THEME_MODE = "System"
THEME_MODE_OPTIONS = ("System", "Dark", "Light")
SETTINGS_FILE = Path(__file__).with_name("translate_settings.json")
MODEL_PRICING_FILE = Path(__file__).with_name("openai_model_prices.json")

# ── Layout spacing ─────────────────────────────────────────────────────────────
# All paddings and margins are multiples of GS (Grid Size).
# Change GS here to scale all spacing throughout the app uniformly.
GS = 4  # base grid unit in pixels

def _coerce_float(value: object) -> float | None:
	try:
		return None if value is None else float(value)
	except Exception:
		return None

def _canon_lang(value: str) -> str:
	raw = str(value or "").strip().replace("_", "-")
	if not raw or not re.fullmatch(r"[A-Za-z]{2,3}(?:-[A-Za-z0-9]{2,8})*", raw):
		return ""
	parts = [p for p in raw.split("-") if p]
	if not parts:
		return ""
	out = [parts[0].lower()]
	for subtag in parts[1:]:
		if len(subtag) == 2 and subtag.isalpha():
			out.append(subtag.upper())
		elif len(subtag) == 4 and subtag.isalpha():
			out.append(subtag.title())
		else:
			out.append(subtag.lower())
	return "-".join(out)

def _clean_language_codes(values: object) -> list[str]:
	items = values if isinstance(values, list) else DEFAULT_LANGUAGE_CODES
	out: list[str] = []
	seen: set[str] = set()
	for value in items:
		code = _canon_lang(str(value or ""))
		if code and code not in seen:
			seen.add(code)
			out.append(code)
	return out or list(DEFAULT_LANGUAGE_CODES)

def _normalize_theme_mode(value: object) -> str:
	mode = str(value or "").strip().lower()
	if mode == "dark":
		return "Dark"
	if mode == "light":
		return "Light"
	return DEFAULT_THEME_MODE

def _resolve_sv_ttk_theme(theme_mode: str) -> str:
	mode = _normalize_theme_mode(theme_mode)
	if mode == "Dark":
		return "dark"
	if mode == "Light":
		return "light"
	detected = str(darkdetect.theme() or "").strip().lower()
	return "dark" if detected.startswith("dark") else "light"

def load_model_pricing_document() -> dict[str, Any]:
	if not MODEL_PRICING_FILE.exists():
		return {}
	try:
		content = MODEL_PRICING_FILE.read_text(encoding="utf-8")
		data = json.loads(content)
		return data if isinstance(data, dict) else {}
	except Exception:
		return {}

def _should_omit_temperature(model_id: str) -> bool:
	"""Check if a model should not receive the temperature parameter.
	gpt-5, o1, and o3 models do not support temperature or restrict it to default value only.
	"""
	model_lower = str(model_id or "").strip().lower()
	for sep in ("/", ":"):
		if sep in model_lower:
			model_lower = model_lower.split(sep)[-1].strip()
	return model_lower.startswith(("gpt-5", "o1", "o3"))

def _supports_reasoning_effort(model_id: str) -> bool:
	"""Check if a model accepts the reasoning_effort parameter."""
	model_lower = str(model_id or "").strip().lower()
	for sep in ("/", ":"):
		if sep in model_lower:
			model_lower = model_lower.split(sep)[-1].strip()
	return model_lower.startswith(("gpt-5", "o1", "o3", "o4"))

def load_model_pricing_map() -> dict[str, dict[str, float | None]]:
	data = load_model_pricing_document()
	try:
		models = data.get("models") if isinstance(data, dict) else None
		if not isinstance(models, dict):
			return {}

		pricing_map: dict[str, dict[str, float | None]] = {}
		for model_name, raw_pricing in models.items():
			if not isinstance(model_name, str) or not isinstance(raw_pricing, dict):
				continue
			input_usd = _coerce_float(raw_pricing.get("input_usd"))
			output_usd = _coerce_float(raw_pricing.get("output_usd"))
			if input_usd is None or output_usd is None:
				continue
			pricing_map[model_name.strip().lower()] = {
				"input": input_usd,
				"cached_input": _coerce_float(raw_pricing.get("cached_input_usd")),
				"output": output_usd,
			}
		return pricing_map
	except Exception:
		return {}

def load_eur_conversion_rate() -> float | None:
	doc = load_model_pricing_document()
	meta = doc.get("_meta") if isinstance(doc.get("_meta"), dict) else {}
	rate = _coerce_float(meta.get("eur_conversion_rate"))
	if rate is None or rate <= 0:
		return None
	return float(rate)

MODEL_PRICING_USD_PER_1M = load_model_pricing_map()

TRANSLATION_SYSTEM_PROMPT = (
	"You are a DOCX translation engine. Output strict JSON object only with key 'translated'. "
	"Preserve all anchor tokens exactly and in order: [[MATH_xxxx]], [[OBJ_xxxx]], "
	"[[CLR_xxxx]]...[[/CLR_xxxx]], and, when present, [[PARA_xxxx]]...[[/PARA_xxxx]] and "
	"[[CHST_xxxx]]...[[/CHST_xxxx]]. Do not rename, remove, duplicate, or reorder anchors. "
	"Keep text inside anchors translatable, keep anchor wrappers exact. Preserve **bold** markers "
	"semantically in translated output using **...** only. No markdown, comments, or extra keys."
)

def compact_json(value: object) -> str:
	return json.dumps(value, ensure_ascii=False, separators=(",", ":"), sort_keys=True)

def normalize_runtime_glossary_entries(entries: list[dict[str, str]] | None) -> list[dict[str, str]]:
	if not entries:
		return []
	seen: set[tuple[str, str]] = set()
	normalized: list[dict[str, str]] = []
	for item in entries:
		if not isinstance(item, dict):
			continue
		left = str(item.get("left", "")).strip()
		right = str(item.get("right", "")).strip()
		if not left or not right:
			continue
		key = (left.lower(), right.lower())
		if key in seen:
			continue
		seen.add(key)
		normalized.append({"left": left, "right": right})
	normalized.sort(key=lambda x: (x["left"].lower(), x["right"].lower()))
	return normalized

def filter_glossary_for_text(
	entries: list[dict[str, str]],
	text: str,
) -> list[dict[str, str]]:
	if not entries or not text:
		return []
	text_lower = text.lower()
	return [
		item for item in entries
		if item["left"].lower() in text_lower or item["right"].lower() in text_lower
	]

_DEFAULT_CHAR_STYLE_IDS = frozenset({
	"defaultparagraphfont",
	"absatz-standardschriftart",
})
_DEFAULT_CHAR_STYLE_NAMES = frozenset({
	"default paragraph font",
	"absatz-standardschriftart",
	"paragraph default font",
})

def _norm_style_key(value: str) -> str:
	return re.sub(r"\s+", " ", str(value or "").strip().lower())

def extract_model_name(model_value: str) -> str:
	text = str(model_value or "").strip()
	if " | In: " in text:
		text = text.split(" | In: ", 1)[0].strip()
	if " (In: " in text:
		text = text.split(" (In: ", 1)[0].strip()
	return text

def normalize_model_name(model_value: str) -> str:
	name = extract_model_name(model_value).lower()
	for sep in ("/", ":"):
		if sep in name:
			name = name.split(sep)[-1].strip()
	return name

def round_money_2(value: float | None) -> float | None:
	if value is None:
		return None
	return round(float(value), 2)

def _format_price_pair(usd_value: float | None, eur_rate: float | None) -> str:
	if usd_value is None:
		return "-$ | -€"
	usd = round_money_2(float(usd_value))
	if usd is None:
		return "-$ | -€"
	if eur_rate is not None:
		eur = round_money_2(float(usd) * eur_rate)
		if eur is not None:
			return f"{float(usd):.2f}$ | {float(eur):.2f}€"
	return f"{float(usd):.2f}$ | -€"

def model_price_texts(model_name: str) -> tuple[str, str, str]:
	model_id = extract_model_name(model_name)
	pricing = model_pricing_for(model_id)
	if pricing is None:
		return "-$ | -€", "-$ | -€", "-$ | -€"
	input_usd = _coerce_float(pricing.get("input"))
	cached_input_usd = _coerce_float(pricing.get("cached_input"))
	output_usd = _coerce_float(pricing.get("output"))
	eur_rate = load_eur_conversion_rate()
	return (
		_format_price_pair(input_usd, eur_rate),
		_format_price_pair(cached_input_usd, eur_rate),
		_format_price_pair(output_usd, eur_rate),
	)

def build_model_choice_label(model_name: str) -> str:
	model_id = extract_model_name(model_name)
	input_text, cached_text, output_text = model_price_texts(model_id)
	return f"{model_id} | In: {input_text} | Cached: {cached_text} | Out: {output_text}"

def is_default_char_style(style_name: str, style_id: str) -> bool:
	name_key = _norm_style_key(style_name)
	id_key = _norm_style_key(style_id)
	return id_key in _DEFAULT_CHAR_STYLE_IDS or name_key in _DEFAULT_CHAR_STYLE_NAMES

class HardAnchorFallbackWarning(Exception):
	def __init__(self, message: str, fallback_text: str, diagnostics: list[str] | None = None) -> None:
		super().__init__(message)
		self.fallback_text = fallback_text
		self.diagnostics = diagnostics or []

def _anchor_context(text: str, pos: int, window: int = 24) -> str:
	if pos < 0:
		return "-"
	start = max(0, pos - window)
	end = min(len(text), pos + window)
	ctx = text[start:end].replace("\n", " ")
	ctx = re.sub(r"\s+", " ", ctx).strip()
	return ctx if ctx else "-"

def build_hard_anchor_diagnostics(
	source_tagged_text: str,
	output_text: str,
	expected_tokens: list[str],
) -> list[str]:
	lines: list[str] = []
	out_matches = [(m.group(0), m.start()) for m in re.finditer(r"\[\[(?:MATH|OBJ)_\d{4}\]\]", output_text)]
	out_tokens = [token for token, _ in out_matches]

	def _ctx_short(value: str, limit: int = 42) -> str:
		clean = re.sub(r"\s+", " ", (value or "").strip())
		if len(clean) <= limit:
			return clean
		return f"{clean[:limit - 3]}..."

	lines.append(f"Expected hard anchors: {len(expected_tokens)} | Output hard anchors: {len(out_tokens)}")
	lines.append("Input anchors (expected)")
	lines.append("  #  Token         Type  SrcPos  Context")
	lines.append("  --  ------------  ----  ------  ------------------------------------------")
	for idx, token in enumerate(expected_tokens, start=1):
		src_pos = source_tagged_text.find(token)
		type_name = token[2: token.find("_")] if "_" in token else "?"
		src_ctx = _ctx_short(_anchor_context(source_tagged_text, src_pos))
		lines.append(f"  {idx:>2}  {token:<12}  {type_name:<4}  {src_pos:>6}  {src_ctx}")

	lines.append("Output anchors (actual)")
	lines.append("  #  Token         Type  OutPos  Context")
	lines.append("  --  ------------  ----  ------  ------------------------------------------")
	if out_matches:
		for idx, (token, out_pos) in enumerate(out_matches, start=1):
			type_name = token[2: token.find("_")] if "_" in token else "?"
			out_ctx = _ctx_short(_anchor_context(output_text, out_pos))
			lines.append(f"  {idx:>2}  {token:<12}  {type_name:<4}  {out_pos:>6}  {out_ctx}")
	else:
		lines.append("  - (no hard anchors found in output)")

	return lines

def replace_many_exact(text: str, replacements: dict[str, str]) -> str:
	if not replacements:
		return text
	out = text
	temp_map: dict[str, str] = {}
	for idx, (old, new) in enumerate(replacements.items(), start=1):
		if old == new:
			continue
		temp = f"@@TMP_REMAP_{idx:04d}@@"
		while temp in out:
			temp = f"{temp}_X"
		temp_map[temp] = new
		out = out.replace(old, temp)

	for temp, new in temp_map.items():
		out = out.replace(temp, new)
	return out

def is_docx_open_in_word(input_path: Path) -> bool:
	try:
		lock_file = input_path.with_name(f"~${input_path.name}")
		return lock_file.exists()
	except Exception:
		return False

def iter_paragraphs(container) -> Iterable:
	for paragraph in container.paragraphs:
		yield paragraph
	for table in container.tables:
		for row in table.rows:
			for cell in row.cells:
				if isinstance(cell, _Cell):
					yield from iter_paragraphs(cell)

def collect_all_paragraphs(doc) -> list:
	seen = set()
	out = []

	def add_paragraphs(paragraph_iterable: Iterable) -> None:
		for paragraph in paragraph_iterable:
			node_id = id(paragraph._p)
			if node_id in seen:
				continue
			seen.add(node_id)
			out.append(paragraph)

	add_paragraphs(iter_paragraphs(doc))
	for section in doc.sections:
		add_paragraphs(iter_paragraphs(section.header))
		add_paragraphs(iter_paragraphs(section.footer))
		add_paragraphs(iter_paragraphs(section.first_page_header))
		add_paragraphs(iter_paragraphs(section.first_page_footer))
		add_paragraphs(iter_paragraphs(section.even_page_header))
		add_paragraphs(iter_paragraphs(section.even_page_footer))
	return out

def get_text_nodes(paragraph) -> list:
	text_nodes = []
	w_t = qn("w:t")

	for run in paragraph.runs:
		for child in run._r.iterchildren():
			if child.tag == w_t:
				text_nodes.append(child)
	return text_nodes

def sanitize_model_json(raw_text: str) -> dict:
	text = raw_text.strip()
	if text.startswith("```"):
		text = re.sub(r"^```[a-zA-Z0-9_-]*\\s*", "", text)
		text = re.sub(r"\\s*```$", "", text)
	return json.loads(text)

def is_non_default_color_node(color_node) -> bool:
	if color_node is None:
		return False
	val = str(color_node.get(qn("w:val")) or "").strip().lower()
	theme = str(color_node.get(qn("w:themeColor")) or "").strip().lower()
	tint = str(color_node.get(qn("w:themeTint")) or "").strip().lower()
	shade = str(color_node.get(qn("w:themeShade")) or "").strip().lower()
	if not val and not theme:
		return False
	if val in {"auto", "000000", "00000000"} and not theme:
		return False
	# Word default text often appears as themeColor="text1". Treat as default unless altered.
	if theme in {"text1", "background1"} and val in {"", "auto", "000000", "00000000"} and not tint and not shade:
		return False
	return True

def describe_color_node(color_node) -> str:
	if color_node is None:
		return "-"
	val = str(color_node.get(qn("w:val")) or "-")
	theme = str(color_node.get(qn("w:themeColor")) or "-")
	tint = str(color_node.get(qn("w:themeTint")) or "-")
	shade = str(color_node.get(qn("w:themeShade")) or "-")
	return f"val={val}, theme={theme}, tint={tint}, shade={shade}"

def build_math_tagged_paragraph(
	paragraph,
	selected_char_styles: set[str] | None = None,
) -> tuple[str, list[str], dict[str, object], list[str], dict[str, object], list[str], dict[str, object], dict[str, object]]:
	parts: list[str] = []
	hard_tokens: list[str] = []
	hard_elements: dict[str, dict[str, object]] = {}
	color_tokens: list[str] = []
	color_styles: dict[str, object] = {}
	char_style_tokens: list[str] = []
	char_style_map: dict[str, object] = {}
	source_bold_runs = 0
	source_bold_segments = 0
	source_named_bold_runs = 0
	source_colored_segments = 0
	used_run_style_labels: set[str] = set()
	used_char_style_labels: set[str] = set()
	used_color_labels: set[str] = set()
	math_idx = 0
	obj_idx = 0
	clr_idx = 0
	chst_idx = 0
	w_r = qn("w:r")
	w_ppr = qn("w:pPr")
	w_t = qn("w:t")
	w_tab = qn("w:tab")
	w_br = qn("w:br")
	w_rpr = qn("w:rPr")
	m_o_math = qn("m:oMath")
	m_o_math_para = qn("m:oMathPara")
	run_by_rid = {id(run._r): run for run in paragraph.runs}

	for p_child in paragraph._p.iterchildren():
		if p_child.tag == w_ppr:
			continue

		if p_child.tag in {m_o_math, m_o_math_para}:
			math_idx += 1
			token = f"[[MATH_{math_idx:04d}]]"
			parts.append(token)
			hard_tokens.append(token)
			hard_elements[token] = {"scope": "paragraph", "element": deepcopy(p_child)}
			continue

		if p_child.tag != w_r:
			obj_idx += 1
			token = f"[[OBJ_{obj_idx:04d}]]"
			parts.append(token)
			hard_tokens.append(token)
			hard_elements[token] = {"scope": "paragraph", "element": deepcopy(p_child)}
			continue

		run_obj = run_by_rid.get(id(p_child))
		run_is_bold = False
		run_color_xml = None
		run_has_char_style = False
		run_char_style_obj = None
		if run_obj is not None:
			style_name_raw = (run_obj.style.name or "").strip() if run_obj.style is not None else ""
			style_name = style_name_raw.lower()
			style_id_raw = str(getattr(run_obj.style, "style_id", "") or "") if run_obj.style is not None else ""
			style_id = style_id_raw.strip().lower()
			style_label = f"{style_name_raw or '-'}<{style_id_raw or '-'}>"
			if run_obj.style is not None:
				used_run_style_labels.add(style_label)
				if run_obj.style.type == 2:
					used_char_style_labels.add(style_label)
			named_bold = style_name in {"fett", "bold", "strong"} or style_id in {"fett", "bold", "strong"}
			style_bold = bool(getattr(run_obj.style.font, "bold", False)) if run_obj.style is not None else False
			run_is_bold = bool(run_obj.bold) or bool(run_obj.font.bold) or style_bold or named_bold
			if run_is_bold:
				source_bold_runs += 1
			if named_bold:
				source_named_bold_runs += 1

			if (
				run_obj.style is not None
				and run_obj.style.type == 2
				and not named_bold
				and not is_default_char_style(style_name_raw, style_id_raw)
				and (selected_char_styles is not None and style_label in selected_char_styles)
			):
				run_has_char_style = True
				run_char_style_obj = run_obj.style
			else:
				rpr = p_child.find(w_rpr)
				if rpr is not None:
					w_color = qn("w:color")
					color_node = rpr.find(w_color)
					if is_non_default_color_node(color_node):
						run_color_xml = deepcopy(color_node)
						used_color_labels.add(describe_color_node(color_node))

		for child in p_child.iterchildren():
			if child.tag == w_t:
				text = child.text or ""
				content = text
				if run_has_char_style and text.strip():
					chst_idx += 1
					chst_token = f"CHST_{chst_idx:04d}"
					char_style_tokens.append(chst_token)
					char_style_map[chst_token] = run_char_style_obj
					content = f"[[{chst_token}]]{text}[[/{chst_token}]]"
				else:
					if run_is_bold and text.strip():
						content = f"**{content}**"
						source_bold_segments += 1
					if run_color_xml is not None and text.strip():
						clr_idx += 1
						token = f"CLR_{clr_idx:04d}"
						color_tokens.append(token)
						color_styles[token] = deepcopy(run_color_xml)
						content = f"[[{token}]]{content}[[/{token}]]"
						source_colored_segments += 1
				parts.append(content)
			elif child.tag in {m_o_math, m_o_math_para}:
				math_idx += 1
				token = f"[[MATH_{math_idx:04d}]]"
				parts.append(token)
				hard_tokens.append(token)
				hard_elements[token] = {"scope": "run", "element": deepcopy(child)}
			elif child.tag == w_tab:
				parts.append("\t")
			elif child.tag == w_br:
				parts.append("\n")
			elif child.tag != w_rpr:
				obj_idx += 1
				token = f"[[OBJ_{obj_idx:04d}]]"
				parts.append(token)
				hard_tokens.append(token)
				hard_elements[token] = {"scope": "run", "element": deepcopy(child)}

	debug_info = {
		"source_bold_runs": source_bold_runs,
		"source_named_bold_runs": source_named_bold_runs,
		"source_bold_segments": source_bold_segments,
		"source_colored_segments": source_colored_segments,
		"source_has_bold": source_bold_segments > 0,
		"source_has_named_bold": source_named_bold_runs > 0,
		"used_run_styles": sorted(used_run_style_labels),
		"used_char_styles": sorted(used_char_style_labels),
		"used_color_values": sorted(used_color_labels),
	}
	return "".join(parts), hard_tokens, hard_elements, color_tokens, color_styles, char_style_tokens, char_style_map, debug_info

def validate_math_anchor_output(
	translated_text: str,
	hard_tokens: list[str],
	color_tokens: list[str],
	char_style_tokens: list[str] | None = None,
	char_style_meta: dict[str, dict[str, object]] | None = None,
) -> str:
	found_tokens = re.findall(r"\[\[(?:MATH|OBJ)_\d{4}\]\]", translated_text)
	if found_tokens != hard_tokens:
		raise ValueError("Model output has invalid hard-anchor sequence.")

	open_color_tokens = [token[2:-2] for token in re.findall(r"\[\[CLR_\d{4}\]\]", translated_text)]
	close_color_tokens = [token[3:-2] for token in re.findall(r"\[\[/CLR_\d{4}\]\]", translated_text)]
	if open_color_tokens != color_tokens or close_color_tokens != color_tokens:
		raise ValueError("Model output has invalid color-anchor sequence.")

	for token in color_tokens:
		pattern = rf"\[\[{token}\]\](.*?)\[\[/{token}\]\]"
		if not re.search(pattern, translated_text, flags=re.DOTALL):
			raise ValueError("Model output has invalid color-anchor mapping.")

	if char_style_tokens:
		open_chst = [t[2:-2] for t in re.findall(r"\[\[CHST_\d{4}\]\]", translated_text)]
		close_chst = [t[3:-2] for t in re.findall(r"\[\[/CHST_\d{4}\]\]", translated_text)]
		if open_chst != char_style_tokens or close_chst != char_style_tokens:
			meta = char_style_meta or {}

			expected_parts: list[str] = []
			for idx, token in enumerate(char_style_tokens, start=1):
				item = meta.get(token, {})
				style_label = str(item.get("style", "-"))
				src_pos = item.get("src_pos", "-")
				src_text = str(item.get("text", "")).replace("\n", " ")
				expected_parts.append(f"{idx}:{token} style={style_label} src_pos={src_pos} text='{src_text}'")

			open_matches = [(m.group(1), m.start()) for m in re.finditer(r"\[\[(CHST_\d{4})\]\]", translated_text)]
			close_matches = [(m.group(1), m.start()) for m in re.finditer(r"\[\[/CHST_(\d{4})\]\]", translated_text)]

			open_parts: list[str] = []
			for idx, (token, pos) in enumerate(open_matches, start=1):
				item = meta.get(token, {})
				style_label = str(item.get("style", "?"))
				src_text = str(item.get("text", "")).replace("\n", " ")
				open_parts.append(f"{idx}:{token} out_pos={pos} style={style_label} text='{src_text}'")

			close_parts: list[str] = []
			for idx, (num, pos) in enumerate(close_matches, start=1):
				token = f"CHST_{num}"
				item = meta.get(token, {})
				style_label = str(item.get("style", "?"))
				src_text = str(item.get("text", "")).replace("\n", " ")
				close_parts.append(f"{idx}:{token} out_pos={pos} style={style_label} text='{src_text}'")

			raise ValueError(
				"Model output has invalid char-style-anchor sequence. "
				f"ExpectedDetail: [{' ; '.join(expected_parts)}] | "
				f"OutputOpenDetail: [{' ; '.join(open_parts)}] | "
				f"OutputCloseDetail: [{' ; '.join(close_parts)}]"
			)

	return translated_text

def style_object_label(style_obj: object | None) -> str:
	if style_obj is None:
		return "-"
	name = str(getattr(style_obj, "name", "") or "-")
	style_id = str(getattr(style_obj, "style_id", "") or "-")
	return f"{name}<{style_id}>"

def build_char_style_anchor_meta(
	tagged_paragraph: str,
	char_style_tokens: list[str] | None,
	char_style_map: dict[str, object] | None,
) -> dict[str, dict[str, object]]:
	if not char_style_tokens:
		return {}
	meta: dict[str, dict[str, object]] = {}
	pattern = re.compile(r"\[\[(CHST_\d{4})\]\](.*?)\[\[/\1\]\]", flags=re.DOTALL)
	for idx, match in enumerate(pattern.finditer(tagged_paragraph), start=1):
		token = match.group(1)
		src_text = re.sub(r"\s+", " ", match.group(2) or "").strip()
		if len(src_text) > 40:
			src_text = f"{src_text[:37]}..."
		style_obj = (char_style_map or {}).get(token)
		meta[token] = {
			"seq": idx,
			"src_pos": match.start(),
			"text": src_text,
			"style": style_object_label(style_obj),
		}

	for idx, token in enumerate(char_style_tokens, start=1):
		if token not in meta:
			meta[token] = {
				"seq": idx,
				"src_pos": "-",
				"text": "",
				"style": style_object_label((char_style_map or {}).get(token)),
			}

	return meta

def parse_color_anchors(text: str) -> list[tuple[str, str | None]]:
	if not text:
		return []
	pattern = re.compile(r"\[\[(CLR_\d{4})\]\](.*?)\[\[/\1\]\]", flags=re.DOTALL)
	segments: list[tuple[str, str | None]] = []
	last_end = 0
	for match in pattern.finditer(text):
		if match.start() > last_end:
			segments.append((text[last_end:match.start()], None))
		segments.append((match.group(2), match.group(1)))
		last_end = match.end()
	if last_end < len(text):
		segments.append((text[last_end:], None))
	return segments

def parse_char_style_anchors(text: str) -> list[tuple[str, str | None]]:
	if not text:
		return []
	pattern = re.compile(r"\[\[(CHST_\d{4})\]\](.*?)\[\[/\1\]\]", flags=re.DOTALL)
	segments: list[tuple[str, str | None]] = []
	last_end = 0
	for match in pattern.finditer(text):
		if match.start() > last_end:
			segments.append((text[last_end:match.start()], None))
		segments.append((match.group(2), match.group(1)))
		last_end = match.end()
	if last_end < len(text):
		segments.append((text[last_end:], None))
	return segments

def apply_color_to_run(run, color_node: object | None) -> None:
	if color_node is None:
		return
	w_color = qn("w:color")
	rpr = run._r.get_or_add_rPr()
	existing = rpr.find(w_color)
	if existing is not None:
		rpr.remove(existing)
	rpr.append(deepcopy(color_node))

def parse_bold_markdown(text: str) -> list[tuple[str, bool]]:
	if not text:
		return []
	parts = re.split(r"(\*\*.*?\*\*)", text, flags=re.DOTALL)
	segments: list[tuple[str, bool]] = []
	for part in parts:
		if not part:
			continue
		if part.startswith("**") and part.endswith("**") and len(part) >= 4:
			inner = part[2:-2]
			if inner:
				segments.append((inner, True))
		else:
			segments.append((part, False))
	return segments

def _contains_term_ci(text: str, term: str) -> bool:
	if not text or not term:
		return False
	return re.search(re.escape(term), text, flags=re.IGNORECASE) is not None

def _source_term_has_bold_occurrence(tagged_source_text: str, source_term: str) -> bool:
	if not source_term:
		return False
	for segment, is_bold in parse_bold_markdown(tagged_source_text):
		if is_bold and _contains_term_ci(segment, source_term):
			return True
	return False

def _strip_markdown_bold_for_term(text: str, term: str) -> tuple[str, int]:
	if not text or not term:
		return text, 0
	term_pattern = re.sub(r"\\\s+", r"\\s+", re.escape(term.strip()))
	if not term_pattern:
		return text, 0
	pattern = re.compile(rf"\*\*(\s*{term_pattern}\s*)\*\*", flags=re.IGNORECASE)
	result, count = pattern.subn(lambda m: m.group(1), text)
	return result, count

def remove_unwanted_glossary_bold(
	tagged_source_text: str,
	translated_text: str,
	glossary_entries: list[dict[str, str]],
) -> str:
	updated = translated_text
	for item in glossary_entries:
		left = str(item.get("left", "") or "").strip()
		right = str(item.get("right", "") or "").strip()
		if not left or not right:
			continue

		left_in_source = _contains_term_ci(tagged_source_text, left)
		right_in_source = _contains_term_ci(tagged_source_text, right)
		mappings: list[tuple[str, str]] = []
		if left_in_source and not right_in_source:
			mappings.append((left, right))
		elif right_in_source and not left_in_source:
			mappings.append((right, left))

		for source_term, target_term in mappings:
			if _source_term_has_bold_occurrence(tagged_source_text, source_term):
				continue
			updated, _ = _strip_markdown_bold_for_term(updated, target_term)
	return updated

def resolve_preferred_bold_style_name(paragraph) -> tuple[object | None, bool, bool]:
	styles = list(paragraph.part.styles)
	fett_exists = any(str(getattr(s, "style_id", "") or "").strip().lower() == "fett" or (getattr(s, "name", "") or "").strip().lower() == "fett" for s in styles)
	bold_exists = any(str(getattr(s, "style_id", "") or "").strip().lower() == "bold" or (getattr(s, "name", "") or "").strip().lower() == "bold" for s in styles)

	# Prefer character style by style_id, then by name.
	for key in ("fett", "bold", "strong"):
		for style in styles:
			try:
				if style.type != 2:
					continue
			except Exception:
				continue
			style_id = str(getattr(style, "style_id", "") or "").strip().lower()
			style_name = (getattr(style, "name", "") or "").strip().lower()
			if style_id == key or style_name == key:
				return style, fett_exists, bold_exists

	return None, fett_exists, bold_exists

def replace_paragraph_with_translated(
	paragraph,
	translated_text: str,
	hard_tokens: list[str],
	hard_elements: dict[str, dict[str, object]],
	color_styles: dict[str, object],
	char_style_map: dict[str, object] | None = None,
) -> dict[str, int | str | bool]:
	bold_style_obj, fett_exists, bold_exists = resolve_preferred_bold_style_name(paragraph)
	bold_style_label = "-"
	if bold_style_obj is not None:
		bold_style_label = f"{getattr(bold_style_obj, 'name', '-') or '-'}<{getattr(bold_style_obj, 'style_id', '-') or '-'}>"
	output_bold_segments = 0
	applied_by_style = 0
	applied_direct = 0
	p_node = paragraph._p
	w_ppr = qn("w:pPr")
	for child in list(p_node):
		if child.tag != w_ppr:
			p_node.remove(child)

	chunks = re.split(r"(\[\[(?:MATH|OBJ)_\d{4}\]\])", translated_text) if hard_tokens else [translated_text]
	for chunk in chunks:
		if not chunk:
			continue
		if chunk in hard_elements:
			info = hard_elements[chunk]
			element = deepcopy(info["element"])
			if info.get("scope") == "paragraph":
				p_node.append(element)
			else:
				run = paragraph.add_run("")
				run._r.append(element)
			continue
		for chst_text, chst_token in parse_char_style_anchors(chunk):
			chst_style_obj = (char_style_map or {}).get(chst_token) if chst_token else None
			for color_text, color_token in parse_color_anchors(chst_text):
				for segment, is_bold in parse_bold_markdown(color_text):
					if not segment:
						continue
					run = paragraph.add_run(segment)
					if chst_style_obj is not None:
						try:
							run.style = chst_style_obj
						except Exception:
							pass
					elif is_bold:
						output_bold_segments += 1
						if bold_style_obj is not None:
							try:
								run.style = bold_style_obj
								applied_by_style += 1
							except Exception:
								run.bold = True
								applied_direct += 1
						else:
							run.bold = True
							applied_direct += 1
					if color_token:
						apply_color_to_run(run, color_styles.get(color_token))

	return {
		"preferred_style": bold_style_label,
		"style_fett_exists": fett_exists,
		"style_bold_exists": bold_exists,
		"output_bold_segments": output_bold_segments,
		"applied_by_style": applied_by_style,
		"applied_direct": applied_direct,
	}

def refresh_model_pricing_cache() -> None:
	global MODEL_PRICING_USD_PER_1M
	MODEL_PRICING_USD_PER_1M = load_model_pricing_map()

def fetch_usd_to_eur_rate_best_effort() -> tuple[float | None, str | None]:
	for endpoint in FX_RATE_ENDPOINTS:
		payload = _fetch_json_from_endpoint(endpoint)
		if not isinstance(payload, dict):
			continue
		rate: float | None = None
		if endpoint.endswith("/latest/USD"):
			rates = payload.get("rates") if isinstance(payload.get("rates"), dict) else {}
			rate = _coerce_float(rates.get("EUR"))
		else:
			rates = payload.get("rates") if isinstance(payload.get("rates"), dict) else {}
			rate = _coerce_float(rates.get("EUR"))
		if rate is not None and rate > 0:
			return float(rate), endpoint
	return None, None

def _object_to_plain_dict(value: object) -> dict[str, Any]:
	if isinstance(value, dict):
		return value
	for method_name in ("model_dump", "to_dict"):
		method = getattr(value, method_name, None)
		if callable(method):
			try:
				result = method()
				if isinstance(result, dict):
					return result
			except Exception:
				pass
	method_json = getattr(value, "model_dump_json", None)
	if callable(method_json):
		try:
			result = json.loads(method_json())
			if isinstance(result, dict):
				return result
		except Exception:
			pass
	try:
		return dict(getattr(value, "__dict__", {}) or {})
	except Exception:
		return {}

def _coerce_price_value(raw: object) -> float | None:
	if raw is None:
		return None
	if isinstance(raw, (int, float)):
		value = float(raw)
		if math.isfinite(value) and value >= 0:
			return value
		return None
	if isinstance(raw, str):
		match = re.search(r"-?\d+(?:\.\d+)?", raw.replace(",", "."))
		if not match:
			return None
		try:
			value = float(match.group(0))
			if math.isfinite(value) and value >= 0:
				return value
		except Exception:
			return None
	return None

def _iter_numeric_fields(value: object, path: str = "") -> Iterable[tuple[str, float]]:
	if isinstance(value, dict):
		for key, sub in value.items():
			child_path = f"{path}.{key}" if path else str(key)
			yield from _iter_numeric_fields(sub, child_path)
		return
	if isinstance(value, list):
		for idx, sub in enumerate(value):
			child_path = f"{path}[{idx}]"
			yield from _iter_numeric_fields(sub, child_path)
		return
	parsed = _coerce_price_value(value)
	if parsed is not None:
		yield path.lower(), parsed

def _normalize_cost_to_per_1m(path: str, value: float) -> float:
	path_lower = path.lower()
	if any(tag in path_lower for tag in ("pricing.prompt", "pricing.completion", "pricing.input", "pricing.output")):
		return value * 1_000_000.0
	if any(tag in path_lower for tag in ("per_token", "/token", "_token")):
		return value * 1_000_000.0
	return value

def _extract_pricing_from_payload(payload: dict[str, Any]) -> dict[str, float | None] | None:
	if not payload:
		return None

	numeric_fields = [(path, _normalize_cost_to_per_1m(path, val)) for path, val in _iter_numeric_fields(payload)]

	input_candidates: list[float] = []
	output_candidates: list[float] = []
	cached_candidates: list[float] = []
	for path, value in numeric_fields:
		if value < 0:
			continue
		is_input = any(tag in path for tag in ("input", "prompt", "inbound"))
		is_output = any(tag in path for tag in ("output", "completion", "generated", "outbound"))
		is_cached = any(tag in path for tag in ("cached", "cache"))
		looks_like_price = any(tag in path for tag in ("price", "pricing", "cost", "rate", "usd", "token"))
		if not looks_like_price:
			continue
		if is_input and is_cached:
			cached_candidates.append(value)
		elif is_input:
			input_candidates.append(value)
		elif is_output:
			output_candidates.append(value)

	if not input_candidates or not output_candidates:
		return None

	input_usd = min(input_candidates)
	output_usd = min(output_candidates)
	cached_usd = min(cached_candidates) if cached_candidates else None
	return {
		"input": input_usd,
		"output": output_usd,
		"cached_input": cached_usd,
	}

def _extract_pricing_from_third_party_payload(payload: object) -> dict[str, dict[str, float | None]]:
	out: dict[str, dict[str, float | None]] = {}

	items: list[dict[str, Any]] = []
	if isinstance(payload, list):
		items = [item for item in payload if isinstance(item, dict)]
	elif isinstance(payload, dict):
		# Support APIs that return a direct {"model-id": {...pricing...}} mapping.
		for model_id, model_payload in payload.items():
			if not isinstance(model_id, str) or not isinstance(model_payload, dict):
				continue
			entry = {"model": model_id}
			entry.update(model_payload)
			items.append(entry)
		for key in ("data", "items", "models", "results"):
			value = payload.get(key)
			if isinstance(value, list):
				items.extend(item for item in value if isinstance(item, dict))
			elif isinstance(value, dict):
				for model_id, model_payload in value.items():
					if isinstance(model_id, str) and isinstance(model_payload, dict):
						entry = {"model": model_id}
						entry.update(model_payload)
						items.append(entry)
		if not items:
			items = [payload]

	for item in items:
		model_id = item.get("model") or item.get("id") or item.get("name") or item.get("model_id")
		if not isinstance(model_id, str) or not model_id.strip():
			continue

		input_usd = (
			_coerce_price_value(item.get("input_usd"))
			or _coerce_price_value(item.get("input"))
			or _coerce_price_value(item.get("prompt"))
			or _coerce_price_value(item.get("prompt_usd"))
			or _coerce_price_value(item.get("input_price"))
			or _coerce_price_value(item.get("input_cost"))
			or _coerce_price_value(item.get("input_cost_usd"))
			or _coerce_price_value(item.get("input_cost_per_1m"))
			or _coerce_price_value(item.get("input_cost_per_token"))
			or _coerce_price_value(item.get("prompt_price"))
			or _coerce_price_value(item.get("prompt_cost"))
			or _coerce_price_value(item.get("prompt_cost_usd"))
		)
		cached_input_usd = (
			_coerce_price_value(item.get("cached_input_usd"))
			or _coerce_price_value(item.get("cached_input"))
			or _coerce_price_value(item.get("cache_input"))
			or _coerce_price_value(item.get("cached_input_price"))
			or _coerce_price_value(item.get("cached_input_cost"))
			or _coerce_price_value(item.get("cache_input_cost"))
			or _coerce_price_value(item.get("cache_creation_input_cost"))
			or _coerce_price_value(item.get("cache_read_input_cost"))
		)
		output_usd = (
			_coerce_price_value(item.get("output_usd"))
			or _coerce_price_value(item.get("output"))
			or _coerce_price_value(item.get("completion"))
			or _coerce_price_value(item.get("completion_usd"))
			or _coerce_price_value(item.get("output_price"))
			or _coerce_price_value(item.get("output_cost"))
			or _coerce_price_value(item.get("output_cost_usd"))
			or _coerce_price_value(item.get("output_cost_per_1m"))
			or _coerce_price_value(item.get("output_cost_per_token"))
			or _coerce_price_value(item.get("completion_price"))
			or _coerce_price_value(item.get("completion_cost"))
			or _coerce_price_value(item.get("completion_cost_usd"))
		)

		pricing = _extract_pricing_from_payload(item)
		if input_usd is None and pricing is not None:
			input_usd = pricing.get("input")
		if cached_input_usd is None and pricing is not None:
			cached_input_usd = pricing.get("cached_input")
		if output_usd is None and pricing is not None:
			output_usd = pricing.get("output")

		if input_usd is None or output_usd is None:
			continue
		out[normalize_model_name(model_id)] = {
			"input": float(input_usd),
			"cached_input": (float(cached_input_usd) if cached_input_usd is not None else None),
			"output": float(output_usd),
		}

	return out

def _resolve_price_for_model(
	price_map: dict[str, dict[str, float | None]],
	model_id: str,
) -> dict[str, float | None] | None:
	normalized = normalize_model_name(model_id)
	if not normalized:
		return None

	exact = price_map.get(normalized)
	if exact is not None:
		return exact

	parts = normalized.split("-")
	for idx in range(len(parts) - 1, 1, -1):
		candidate = "-".join(parts[:idx])
		if candidate in price_map:
			return price_map[candidate]

	prefix_matches = [
		(name, pricing)
		for name, pricing in price_map.items()
		if normalized.startswith(name)
	]
	if prefix_matches:
		prefix_matches.sort(key=lambda item: len(item[0]), reverse=True)
		return prefix_matches[0][1]

	return None

def _fetch_json_from_endpoint(url: str, api_key: str | None = None) -> object | None:
	headers = {"Content-Type": "application/json", "Accept": "application/json"}
	if api_key:
		headers["Authorization"] = f"Bearer {api_key}"
	for attempt in range(3):
		request_obj = urllib_request.Request(url, headers=headers, method="GET")
		try:
			with urllib_request.urlopen(request_obj, timeout=20) as response:
				raw = response.read()
				text = raw.decode("utf-8", errors="replace")
				return json.loads(text)
		except urllib_error.HTTPError as exc:
			if exc.code in (429, 500, 502, 503, 504) and attempt < 2:
				time.sleep(1.5 * (attempt + 1))
				continue
			return None
		except (urllib_error.URLError, TimeoutError, ValueError):
			if attempt < 2:
				time.sleep(1.0 * (attempt + 1))
				continue
			return None
	return None

def fetch_model_prices_best_effort(
	api_key: str,
	model_ids: list[str],
	progress_callback: Callable[[int, int, str], None] | None = None,
) -> tuple[dict[str, dict[str, float | None]], list[str]]:
	prices: dict[str, dict[str, float | None]] = {}
	diagnostics: list[str] = []
	normalized_targets = {normalize_model_name(mid) for mid in model_ids if normalize_model_name(mid)}
	total_steps = len(THIRD_PARTY_PRICING_API_ENDPOINTS) + len(model_ids)
	completed_steps = 0

	def emit(message: str) -> None:
		if progress_callback is not None:
			progress_callback(completed_steps, total_steps, message)

	emit("Checking pricing APIs...")
	if not normalized_targets:
		return prices, diagnostics

	emit("Checking external pricing APIs...")
	for index, endpoint in enumerate(THIRD_PARTY_PRICING_API_ENDPOINTS, start=1):
		payload_obj = _fetch_json_from_endpoint(endpoint)
		completed_steps += 1
		if payload_obj is None:
			emit(f"External API {index}/{len(THIRD_PARTY_PRICING_API_ENDPOINTS)} not available")
			continue
		ext_prices = _extract_pricing_from_third_party_payload(payload_obj)
		for model_name, pricing in ext_prices.items():
			if pricing.get("input") is not None and pricing.get("output") is not None:
				prices[model_name] = pricing
		if ext_prices:
			diagnostics.append(f"External pricing API returned {len(ext_prices)} model entries: {endpoint}")
		emit(f"External API {index}/{len(THIRD_PARTY_PRICING_API_ENDPOINTS)} checked")

	client = OpenAI(api_key=api_key)
	for index, model_id in enumerate(model_ids, start=1):
		normalized_id = normalize_model_name(model_id)
		completed_steps += 1
		if not normalized_id or _resolve_price_for_model(prices, normalized_id) is not None:
			emit(f"Model {index}/{len(model_ids)} already resolved")
			continue
		try:
			model_payload = _object_to_plain_dict(client.models.retrieve(model_id))
		except Exception:
			emit(f"Model {index}/{len(model_ids)} unavailable")
			continue
		pricing = _extract_pricing_from_payload(model_payload)
		if pricing is not None and pricing.get("input") is not None and pricing.get("output") is not None:
			prices[normalized_id] = pricing
			emit(f"Model {index}/{len(model_ids)} resolved")
		else:
			emit(f"Model {index}/{len(model_ids)} has no pricing payload")

	return prices, diagnostics

def update_model_pricing_file(
	api_key: str,
	model_ids: list[str],
	progress_callback: Callable[[int, int, str], None] | None = None,
) -> dict[str, object]:
	def emit(step: int, total: int, message: str) -> None:
		if progress_callback is not None:
			progress_callback(step, total, message)

	if MODEL_PRICING_FILE.exists():
		try:
			existing_doc = json.loads(MODEL_PRICING_FILE.read_text(encoding="utf-8"))
		except Exception:
			existing_doc = {}
	else:
		existing_doc = {}

	if not isinstance(existing_doc, dict):
		existing_doc = {}
	meta = existing_doc.get("_meta") if isinstance(existing_doc.get("_meta"), dict) else {}
	models_doc = existing_doc.get("models") if isinstance(existing_doc.get("models"), dict) else {}

	unique_models = sorted({extract_model_name(mid) for mid in model_ids if extract_model_name(mid)})
	if not unique_models:
		unique_models = fetch_accessible_models(api_key)
	fetch_steps = len(THIRD_PARTY_PRICING_API_ENDPOINTS) + len(unique_models)
	total_steps = max(fetch_steps + 3, 1)
	emit(0, total_steps, "Preparing model price update...")
	fetched_prices, diagnostics = fetch_model_prices_best_effort(
		api_key,
		unique_models,
		progress_callback=lambda done, total, message: emit(min(1 + done, total_steps - 2), total_steps, message),
	)
	emit(total_steps - 2, total_steps, "Merging fetched prices...")

	updated_count = 0
	resolved_count = 0
	seeded_count = 0
	unresolved: list[str] = []
	for model_id in unique_models:
		fetched = _resolve_price_for_model(fetched_prices, model_id)
		existing_row = models_doc.get(model_id) if isinstance(models_doc.get(model_id), dict) else {}
		if fetched is None or fetched.get("input") is None or fetched.get("output") is None:
			if not existing_row:
				models_doc[model_id] = {
					"family": "unresolved",
					"input_usd": None,
					"cached_input_usd": None,
					"output_usd": None,
				}
				updated_count += 1
				seeded_count += 1
			unresolved.append(model_id)
			continue

		resolved_count += 1
		new_row = {
			"family": str(existing_row.get("family") or "updated"),
			"input_usd": round_money_2(float(fetched["input"])),
			"cached_input_usd": (round_money_2(float(fetched["cached_input"])) if fetched.get("cached_input") is not None else None),
			"output_usd": round_money_2(float(fetched["output"])),
		}
		if existing_row != new_row:
			updated_count += 1
		models_doc[model_id] = new_row

	meta["source"] = "Pricing APIs (best effort)"
	meta["captured_on"] = datetime.now().strftime("%Y-%m-%d")
	meta["currency"] = "USD per 1M tokens"
	fx_rate, fx_source = fetch_usd_to_eur_rate_best_effort()
	if fx_rate is not None:
		meta["eur_conversion_rate"] = fx_rate
	if fx_source:
		meta["eur_source"] = fx_source
	meta["notes"] = [
		"Best-effort refresh from available OpenAI endpoints.",
		"Unresolved models are written with null pricing values.",
	]

	out_doc = {
		"_meta": meta,
		"models": dict(sorted(models_doc.items(), key=lambda item: str(item[0]).lower())),
	}
	MODEL_PRICING_FILE.parent.mkdir(parents=True, exist_ok=True)
	MODEL_PRICING_FILE.write_text(json.dumps(out_doc, ensure_ascii=False, indent=2), encoding="utf-8")
	emit(total_steps, total_steps, "Model price update finished.")

	return {
		"updated_count": updated_count,
		"seeded_count": seeded_count,
		"resolved_count": resolved_count,
		"total_count": len(unique_models),
		"model_ids": unique_models,
		"unresolved_models": unresolved,
		"diagnostics": diagnostics,
	}

def apply_translated_with_appended_hard_elements(
	paragraph,
	translated_plain: str,
	hard_elements: dict[str, dict[str, object]],
	color_styles: dict[str, object],
	char_style_map: dict[str, object] | None = None,
) -> None:
	clean = re.sub(r"\[\[/?(?:MATH|OBJ|CLR|CHST|PARA)_\d{4}\]\]", "", translated_plain)
	clean = re.sub(r"\*\*(.*?)\*\*", r"\1", clean, flags=re.DOTALL)
	clean = re.sub(r"\s+", " ", clean).strip()
	replace_paragraph_with_translated(
		paragraph=paragraph,
		translated_text=clean,
		hard_tokens=[],
		hard_elements={},
		color_styles=color_styles,
		char_style_map=char_style_map,
	)
	p_node = paragraph._p
	for _token, info in hard_elements.items():
		element = deepcopy(info["element"])
		if info.get("scope") == "paragraph":
			p_node.append(element)
		else:
			run = paragraph.add_run(" ")
			run._r.append(element)

def model_pricing_for(model_name: str) -> dict[str, float] | None:
	name = normalize_model_name(model_name)
	if not name:
		return None

	if name in MODEL_PRICING_USD_PER_1M:
		return MODEL_PRICING_USD_PER_1M[name]
	for known, pricing in MODEL_PRICING_USD_PER_1M.items():
		if name.startswith(known):
			return pricing
	return None

def estimate_cost_usd(model_name: str, prompt_tokens: int, completion_tokens: int) -> float | None:
	pricing = model_pricing_for(model_name)
	if pricing is None:
		return None
	in_cost = (max(prompt_tokens, 0) / 1_000_000.0) * float(pricing["input"])
	out_cost = (max(completion_tokens, 0) / 1_000_000.0) * float(pricing["output"])
	return in_cost + out_cost

def fetch_accessible_models(api_key: str) -> list[str]:
	client = OpenAI(api_key=api_key)
	response = client.models.list()
	ids = sorted(
		{
			model.id
			for model in response.data
			if isinstance(model.id, str)
			and (
				model.id.startswith("gpt-")
				or model.id.startswith("chatgpt-")
				or model.id.startswith("codex-")
				or model.id.startswith("o1")
				or model.id.startswith("o3")
				or model.id.startswith("o4")
			)
		}
	)
	if not ids:
		raise ValueError("No matching GPT/O models found for this API key.")
	return ids

def load_gui_settings() -> dict:
	payload: dict[str, Any] = {}
	if SETTINGS_FILE.exists():
		try:
			content = SETTINGS_FILE.read_text(encoding="utf-8")
			data = json.loads(content)
			if isinstance(data, dict):
				payload = data
		except Exception:
			payload = {}

	language_codes = _clean_language_codes(payload.get("language_codes"))
	default_source = language_codes[0]
	default_target = language_codes[1] if len(language_codes) > 1 else language_codes[0]
	source = _canon_lang(str(payload.get("source_language") or default_source)) or default_source
	target = _canon_lang(str(payload.get("target_language") or default_target)) or default_target
	if source not in language_codes:
		source = default_source
	if target not in language_codes:
		target = default_target

	payload["language_codes"] = language_codes
	payload["source_language"] = source
	payload["target_language"] = target
	payload["theme_mode"] = _normalize_theme_mode(payload.get("theme_mode"))
	try:
		SETTINGS_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
	except Exception:
		pass
	return payload

def language_pair_key(lang_a: str, lang_b: str) -> str:
	a = _canon_lang(lang_a).lower()
	b = _canon_lang(lang_b).lower()
	return "||".join(sorted([a, b]))

def normalize_glossary_pairs(raw: object) -> dict[str, list[dict[str, str]]]:
	if not isinstance(raw, dict):
		return {}
	out: dict[str, list[dict[str, str]]] = {}
	for key, value in raw.items():
		if not isinstance(key, str) or not isinstance(value, list):
			continue
		if "||" not in key:
			continue
		left_raw, right_raw = key.split("||", 1)
		normalized_key = language_pair_key(left_raw, right_raw)
		if not normalized_key:
			continue
		pairs: list[dict[str, str]] = []
		for item in value:
			if not isinstance(item, dict):
				continue
			left = str(item.get("left", "")).strip()
			right = str(item.get("right", "")).strip()
			if left and right:
				pairs.append({"left": left, "right": right})
		if pairs:
			out.setdefault(normalized_key, []).extend(pairs)
	return out

def get_glossary_for_pair(
	glossary_pairs: dict[str, list[dict[str, str]]],
	source_language: str,
	target_language: str,
) -> list[dict[str, str]]:
	key = language_pair_key(source_language, target_language)
	entries = glossary_pairs.get(key, [])
	out: list[dict[str, str]] = []
	for item in entries:
		left = str(item.get("left", "")).strip()
		right = str(item.get("right", "")).strip()
		if left and right:
			out.append({"left": left, "right": right})
	return out

def save_gui_settings(
	input_file: str,
	model: str,
	selected_styles: list[str] | None = None,
	selected_char_styles: list[str] | None = None,
	source_language: str | None = None,
	target_language: str | None = None,
	theme_mode: str | None = None,
	debug_mode: bool | None = None,
	glossary_pairs: dict[str, list[dict[str, str]]] | None = None,
	api_key: str | None = None,
	language_codes: list[str] | None = None,
) -> None:
	try:
		payload = load_gui_settings()
		if not isinstance(payload, dict):
			payload = {}

		payload["input_file"] = input_file
		payload["model"] = model
		if selected_styles is not None:
			payload["selected_styles"] = selected_styles
		if selected_char_styles is not None:
			payload["selected_char_styles"] = selected_char_styles
		if source_language is not None:
			payload["source_language"] = _canon_lang(source_language)
		if target_language is not None:
			payload["target_language"] = _canon_lang(target_language)
		if theme_mode is not None:
			payload["theme_mode"] = _normalize_theme_mode(theme_mode)
		if debug_mode is not None:
			payload["debug_mode"] = bool(debug_mode)
		if glossary_pairs is not None:
			payload["glossary_pairs"] = glossary_pairs
		if api_key is not None:
			payload["api_key"] = str(api_key)
		if language_codes is not None:
			payload["language_codes"] = _clean_language_codes(language_codes)

		codes = _clean_language_codes(payload.get("language_codes"))
		payload["language_codes"] = codes
		default_source = codes[0]
		default_target = codes[1] if len(codes) > 1 else codes[0]
		source = _canon_lang(str(payload.get("source_language") or default_source)) or default_source
		target = _canon_lang(str(payload.get("target_language") or default_target)) or default_target
		payload["source_language"] = source if source in codes else default_source
		payload["target_language"] = target if target in codes else default_target
		payload["theme_mode"] = _normalize_theme_mode(payload.get("theme_mode"))

		SETTINGS_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
	except Exception:
		# Settings save failures should not block translation.
		pass

class OpenAIParagraphTranslator:
	def __init__(
		self,
		api_key: str,
		model: str,
		source_lang: str,
		target_lang: str,
		glossary_pairs: list[dict[str, str]] | None = None,
		usage_callback=None,
		request_callback=None,
	) -> None:
		self.client = OpenAI(api_key=api_key)
		self.model = model
		self.source_lang = source_lang
		self.target_lang = target_lang
		self.glossary_pairs = normalize_runtime_glossary_entries(glossary_pairs)
		self.usage_callback = usage_callback
		self.request_callback = request_callback
		self.api_calls = 0
		self.prompt_tokens = 0
		self.completion_tokens = 0
		self.total_tokens = 0
		self.translation_cache: dict[str, str] = {}

	def get_usage_summary(self) -> dict[str, object]:
		estimated_usd = estimate_cost_usd(self.model, self.prompt_tokens, self.completion_tokens)
		eur_rate = load_eur_conversion_rate()
		estimated_eur = (estimated_usd * eur_rate) if (estimated_usd is not None and eur_rate is not None) else None
		pricing = model_pricing_for(self.model)
		return {
			"model": self.model,
			"api_calls": int(self.api_calls),
			"prompt_tokens": int(self.prompt_tokens),
			"completion_tokens": int(self.completion_tokens),
			"total_tokens": int(self.total_tokens),
			"estimated_cost_usd": estimated_usd,
			"estimated_cost_eur": estimated_eur,
			"pricing_known": pricing is not None,
			"pricing": pricing,
			"eur_rate": eur_rate,
		}

	def translate_tagged_paragraph(
		self,
		tagged_paragraph: str,
		hard_tokens: list[str],
		color_tokens: list[str],
		char_style_tokens: list[str] | None = None,
		char_style_map: dict[str, object] | None = None,
		attempt_callback=None,
		post_validate=None,
	) -> str:
		last_bad_output: str | None = None
		last_error: str | None = None
		char_style_meta = build_char_style_anchor_meta(tagged_paragraph, char_style_tokens, char_style_map)
		cache_key = json.dumps(
			{
				"tagged_paragraph": tagged_paragraph,
				"hard_tokens": hard_tokens,
				"color_tokens": color_tokens,
				"char_style_tokens": char_style_tokens or [],
			},
			ensure_ascii=False,
			sort_keys=True,
		)
		cached = self.translation_cache.get(cache_key)
		if cached is not None:
			if post_validate:
				post_validate(cached)
			return cached

		for attempt in range(1, 4):
			if attempt_callback:
				attempt_callback(attempt)

			has_para_wrappers = "[[PARA_" in tagged_paragraph

			user_payload = {
				"task": "docx_translate",
				"source_language": self.source_lang,
				"target_language": self.target_lang,
				"preserve_para_wrappers": has_para_wrappers,
				"preserve_char_style_wrappers": bool(char_style_tokens),
				"tagged_paragraph": tagged_paragraph,
				"output_format": {"translated": "string"},
			}
			if attempt > 1 and last_bad_output is not None:
				user_payload["retry"] = {
					"attempt": attempt - 1,
					"last_error": str(last_error or "-"),
					"previous_bad_output": last_bad_output,
					"instruction": "Fix anchor violations and return corrected translation.",
				}

			relevant_glossary = filter_glossary_for_text(self.glossary_pairs, tagged_paragraph)
			messages = [{"role": "system", "content": TRANSLATION_SYSTEM_PROMPT}]
			if relevant_glossary:
				messages.append(
					{
						"role": "system",
						"content": (
							"Terminology table (bidirectional). If either side appears, use the paired term exactly where appropriate. Do not add **bold** just because a glossary term exists; bold must only follow source formatting: "
							+ compact_json(relevant_glossary)
						),
					}
				)
			messages.append({"role": "user", "content": compact_json(user_payload)})

			# Build request kwargs, conditionally including temperature
			request_kwargs = {
				"model": self.model,
				"messages": messages,
				"response_format": {"type": "json_object"},
			}
			if not _should_omit_temperature(self.model):
				request_kwargs["temperature"] = 0
			if _supports_reasoning_effort(self.model):
				request_kwargs["reasoning_effort"] = "low"
			response = self.client.chat.completions.create(**request_kwargs)
			self.api_calls += 1
			if self.request_callback:
				try:
					self.request_callback(
						{
							"api_calls": int(self.api_calls),
							"model": self.model,
							"request_id": str(getattr(response, "_request_id", "") or getattr(response, "request_id", "") or "-"),
							"response_id": str(getattr(response, "id", "") or "-"),
						}
					)
				except Exception:
					pass
			usage = getattr(response, "usage", None)
			if usage is not None:
				try:
					self.prompt_tokens += int(getattr(usage, "prompt_tokens", 0) or 0)
					self.completion_tokens += int(getattr(usage, "completion_tokens", 0) or 0)
					self.total_tokens += int(getattr(usage, "total_tokens", 0) or 0)
				except Exception:
					pass
			if self.usage_callback:
				try:
					self.usage_callback(self.get_usage_summary())
				except Exception:
					pass
			content = response.choices[0].message.content or "{}"
			data = sanitize_model_json(content)
			translated_text = data.get("translated", "")

			if not isinstance(translated_text, str):
				last_bad_output = str(translated_text)
				last_error = "Model output is missing 'translated' text."
				if attempt == 3:
					raise ValueError(last_error)
				continue

			if relevant_glossary:
				translated_text = remove_unwanted_glossary_bold(tagged_paragraph, translated_text, relevant_glossary)

			try:
				validated_text = validate_math_anchor_output(
					translated_text,
					hard_tokens,
					color_tokens,
					char_style_tokens,
					char_style_meta,
				)
				if post_validate:
					post_validate(validated_text)
				self.translation_cache[cache_key] = validated_text
				return validated_text
			except ValueError as exc:
				last_bad_output = translated_text
				last_error = str(exc)
				if attempt == 3:
					if "invalid hard-anchor sequence" in last_error:
						cleaned = re.sub(r"\[\[(?:MATH|OBJ)_\d{4}\]\]", "", last_bad_output)
						diagnostics = build_hard_anchor_diagnostics(tagged_paragraph, last_bad_output, hard_tokens)
						raise HardAnchorFallbackWarning(last_error, cleaned, diagnostics) from exc
					raise

		raise ValueError("Paragraph translation failed after multiple retries.")

def format_openai_exception(exc: Exception) -> tuple[str, list[str]]:
	status_code = getattr(exc, "status_code", None)
	request_id = getattr(exc, "request_id", None)
	body = getattr(exc, "body", None)

	error_message = str(exc)
	error_type = type(exc).__name__
	error_code = None
	error_param = None

	if isinstance(body, dict):
		error_obj = body.get("error") if isinstance(body.get("error"), dict) else body
		error_message = str(error_obj.get("message") or error_message)
		error_type = str(error_obj.get("type") or error_type)
		error_code = error_obj.get("code")
		error_param = error_obj.get("param")

	is_ai_error = (
		status_code is not None
		or bool(request_id)
		or isinstance(body, dict)
		or type(exc).__module__.startswith("openai")
	)

	details = [
		f"Message: {error_message}",
		f"Type: {error_type}",
	]
	if is_ai_error:
		details.extend(
			[
				f"HTTP Status: {status_code if status_code is not None else '-'}",
				f"Code: {error_code if error_code else '-'}",
				f"Param: {error_param if error_param else '-'}",
				f"Request ID: {request_id if request_id else '-'}",
			]
		)
	return error_message, details

def language_code_for_label(language_label: str) -> str:
	return _canon_lang(language_label) or "en-US"

def language_suffix_for_label(language_label: str) -> str:
	return language_code_for_label(language_label).upper()

def apply_document_language(doc: Document, language_label: str) -> None:
	lang_code = language_code_for_label(language_label)
	w_lang = qn("w:lang")

	for paragraph in collect_all_paragraphs(doc):
		for run in paragraph.runs:
			rpr = run._r.get_or_add_rPr()
			lang = rpr.find(w_lang)
			if lang is None:
				lang = OxmlElement("w:lang")
				rpr.append(lang)
			lang.set(qn("w:val"), lang_code)
			lang.set(qn("w:eastAsia"), lang_code)
			lang.set(qn("w:bidi"), lang_code)

def build_output_path(input_path: Path, target_lang: str) -> Path:
	suffix = language_suffix_for_label(target_lang)
	return input_path.with_name(f"{input_path.stem}_{suffix}{input_path.suffix}")

def build_last_source_snapshot_path(input_path: Path, target_lang: str) -> Path:
	suffix = language_suffix_for_label(target_lang)
	return input_path.with_name(f"{input_path.stem}_{suffix}_last{input_path.suffix}")

def is_translatable_paragraph(paragraph, selected_styles: set[str]) -> bool:
	style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
	if style_name not in selected_styles:
		return False

	text_nodes = get_text_nodes(paragraph)
	if not text_nodes:
		return False

	if not any((node.text or "").strip() for node in text_nodes):
		return False

	return True

def find_translatable_paragraph_entries(paragraphs: list, selected_styles: set[str]) -> list[tuple[int, object]]:
	return [
		(idx, paragraph)
		for idx, paragraph in enumerate(paragraphs)
		if is_translatable_paragraph(paragraph, selected_styles)
	]

def find_translatable_paragraphs(paragraphs: list, selected_styles: set[str]) -> list:
	return [paragraph for _, paragraph in find_translatable_paragraph_entries(paragraphs, selected_styles)]

def paragraph_source_signature(paragraph, selected_char_styles: set[str]) -> str:
	style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
	tagged_paragraph, hard_tokens, _, color_tokens, _, char_style_tokens, _, _ = build_math_tagged_paragraph(
		paragraph,
		selected_char_styles=selected_char_styles,
	)
	payload = {
		"style": style_name,
		"tagged": tagged_paragraph,
		"hard": hard_tokens,
		"color": color_tokens,
		"char": char_style_tokens,
	}
	return hashlib.sha256(compact_json(payload).encode("utf-8")).hexdigest()

def replace_paragraph_xml(target_paragraph, source_paragraph) -> None:
	new_element = deepcopy(source_paragraph._p)
	old_element = target_paragraph._p
	parent = old_element.getparent()
	if parent is None:
		return
	parent.replace(old_element, new_element)
	target_paragraph._p = new_element
	target_paragraph._element = new_element

def apply_incremental_translation_seed(
	current_doc: Document,
	output_path: Path,
	snapshot_path: Path,
	selected_styles: set[str],
	selected_char_styles: set[str],
	precheck_progress_callback: Callable[[int, int, str], None] | None = None,
) -> tuple[list, list[int], Document, list] | None:
	def report_precheck(processed: int, total: int, label: str) -> None:
		if precheck_progress_callback:
			try:
				precheck_progress_callback(processed, total, label)
			except Exception:
				pass

	if not output_path.exists() or not snapshot_path.exists():
		return None

	try:
		snapshot_doc = Document(str(snapshot_path))
		existing_output_doc = Document(str(output_path))
	except Exception:
		return None

	current_all = collect_all_paragraphs(current_doc)
	snapshot_all = collect_all_paragraphs(snapshot_doc)
	output_all = collect_all_paragraphs(existing_output_doc)
	if len(snapshot_all) != len(output_all):
		return None

	current_entries = find_translatable_paragraph_entries(current_all, selected_styles)
	snapshot_entries = find_translatable_paragraph_entries(snapshot_all, selected_styles)
	output_entries = find_translatable_paragraph_entries(output_all, selected_styles)
	if [idx for idx, _ in snapshot_entries] != [idx for idx, _ in output_entries]:
		return None

	def build_legacy_style_blocks(entries: list[tuple[int, object]]) -> list[dict[str, object]]:
		blocks: list[dict[str, object]] = []
		current_paragraphs: list[object] = []
		current_indices: list[int] = []
		current_style: str | None = None
		last_index: int | None = None
		for abs_index, paragraph in entries:
			style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
			if not current_paragraphs:
				current_paragraphs = [paragraph]
				current_indices = [abs_index]
				current_style = style_name
				last_index = abs_index
				continue
			if style_name == current_style and last_index is not None and abs_index == last_index + 1:
				current_paragraphs.append(paragraph)
				current_indices.append(abs_index)
				last_index = abs_index
				continue
			blocks.append({"style": current_style or "", "paragraphs": current_paragraphs, "indices": current_indices})
			current_paragraphs = [paragraph]
			current_indices = [abs_index]
			current_style = style_name
			last_index = abs_index
		if current_paragraphs:
			blocks.append({"style": current_style or "", "paragraphs": current_paragraphs, "indices": current_indices})
		return blocks

	def block_signature(block: dict[str, object]) -> str:
		paragraphs = list(block.get("paragraphs", []))
		payload = {
			"style": str(block.get("style", "")),
			"paragraphs": [paragraph_source_signature(paragraph, selected_char_styles) for paragraph in paragraphs],
		}
		return hashlib.sha256(compact_json(payload).encode("utf-8")).hexdigest()

	snapshot_blocks = build_legacy_style_blocks(snapshot_entries)
	current_blocks = build_legacy_style_blocks(current_entries)

	n_blocks = max(1, len(snapshot_blocks))
	report_precheck(0, n_blocks, "")

	output_blocks_by_snapshot: list[list[object]] = []
	output_pos = 0
	for block in snapshot_blocks:
		block_len = len(list(block.get("paragraphs", [])))
		if output_pos + block_len > len(output_entries):
			return None
		output_block = [paragraph for _, paragraph in output_entries[output_pos: output_pos + block_len]]
		output_blocks_by_snapshot.append(output_block)
		output_pos += block_len
	if output_pos != len(output_entries):
		return None

	snapshot_block_signatures: list[str] = []
	for idx, block in enumerate(snapshot_blocks, start=1):
		snapshot_block_signatures.append(block_signature(block))
		report_precheck(idx, n_blocks, "")

	current_block_signatures: list[str] = []
	for block in current_blocks:
		current_block_signatures.append(block_signature(block))

	unchanged_current_block_positions: set[int] = set()
	matcher = SequenceMatcher(a=snapshot_block_signatures, b=current_block_signatures, autojunk=False)
	for snapshot_pos, current_pos, size in matcher.get_matching_blocks():
		if size <= 0:
			continue
		for offset in range(size):
			unchanged_current_block_positions.add(current_pos + offset)

	changed_paragraphs: list = []
	changed_indices: list[int] = []
	for block_pos, block in enumerate(current_blocks):
		if block_pos in unchanged_current_block_positions:
			continue
		changed_paragraphs.extend(list(block.get("paragraphs", [])))
		changed_indices.extend(list(block.get("indices", [])))

	report_precheck(n_blocks, n_blocks, "")

	return changed_paragraphs, changed_indices, existing_output_doc, output_all

def paragraph_label(paragraph, max_chars: int = 12) -> str:
	style_name = (paragraph.style.name or "No Style").strip() if paragraph.style else "No Style"
	text = re.sub(r"\s+", " ", paragraph.text or "").strip()
	if not text:
		preview = "(leer)"
	elif len(text) > max_chars:
		preview = f"{text[:max_chars]}..."
	else:
		preview = text
	return f"{style_name} | {preview}"

def collect_styles_from_docx(input_path: Path) -> list[str]:
	doc = Document(str(input_path))
	styles: set[str] = set()
	for para in collect_all_paragraphs(doc):
		if para.style and para.style.name:
			styles.add(para.style.name.strip())
	return sorted(styles)

def collect_character_styles_from_docx(input_path: Path, include_named_bold_styles: bool = False) -> list[str]:
	doc = Document(str(input_path))
	styles: set[str] = set()
	for para in collect_all_paragraphs(doc):
		for run in para.runs:
			style_obj = run.style
			if style_obj is None:
				continue
			try:
				if style_obj.type != 2:
					continue
			except Exception:
				continue
			style_name_raw = str(getattr(style_obj, "name", "") or "").strip()
			style_id_raw = str(getattr(style_obj, "style_id", "") or "").strip()
			if is_default_char_style(style_name_raw, style_id_raw):
				continue
			style_name = style_name_raw.lower()
			style_id = style_id_raw.lower()
			if (not include_named_bold_styles) and (
				style_name in {"fett", "bold", "strong"} or style_id in {"fett", "bold", "strong"}
			):
				continue
			styles.add(f"{style_name_raw or '-'}<{style_id_raw or '-'}>")
	return sorted(styles)

def collect_all_defined_styles_from_docx(input_path: Path) -> tuple[list[str], list[str]]:
	doc = Document(str(input_path))
	para_styles: list[str] = []
	char_styles: list[str] = []
	for style in doc.styles:
		try:
			if style.type == 1 and style.name:
				para_styles.append(style.name.strip())
			elif style.type == 2:
				name_raw = str(getattr(style, "name", "") or "").strip()
				id_raw = str(getattr(style, "style_id", "") or "").strip()
				if not is_default_char_style(name_raw, id_raw):
					char_styles.append(f"{name_raw or '-'}<{id_raw or '-'}>")
		except Exception:
			continue
	return sorted(para_styles), sorted(char_styles)

def remove_styles_from_docx(input_path: Path, para_style_names: list[str], char_style_labels: list[str]) -> int:
	doc = Document(str(input_path))
	removed = 0
	styles_element = doc.part.styles._element
	for style in list(doc.styles):
		try:
			name = (style.name or "").strip()
			sid = (getattr(style, "style_id", "") or "").strip()
			label = f"{name}<{sid}>"
			if (style.type == 1 and name in para_style_names) or \
			   (style.type == 2 and label in char_style_labels):
				styles_element.remove(style._element)
				removed += 1
		except Exception:
			continue
	doc.save(str(input_path))
	return removed

def _read_style_ui_priority(style_obj) -> int:
	try:
		ui_node = style_obj._element.find(qn("w:uiPriority"))
		if ui_node is None:
			return 9999
		val = str(ui_node.get(qn("w:val")) or "").strip()
		return int(val) if val else 9999
	except Exception:
		return 9999

def _read_latent_style_ui_priority(latent_style_el) -> int:
	try:
		val = str(latent_style_el.get(qn("w:uiPriority")) or "").strip()
		return int(val) if val else 9999
	except Exception:
		return 9999

def collect_all_styles_with_priority_from_docx(input_path: Path) -> list[dict[str, object]]:
	doc = Document(str(input_path))
	styles_root = doc.part.styles._element
	type_labels = {
		1: "Paragraph",
		2: "Character",
		3: "Table",
		4: "List",
	}
	items: list[dict[str, object]] = []
	explicit_name_keys: set[str] = set()
	for style in doc.styles:
		try:
			name = str(getattr(style, "name", "") or "").strip() or "-"
			style_id = str(getattr(style, "style_id", "") or "").strip() or "-"
			style_type = int(getattr(style, "type", 0) or 0)
			explicit_name_keys.add(_norm_style_key(name))
			items.append({
				"key": f"style:{_norm_style_key(style_id)}",
				"name": name,
				"style_id": style_id,
				"type": type_labels.get(style_type, f"Type {style_type}"),
				"priority": _read_style_ui_priority(style),
			})
		except Exception:
			continue
	latent_styles_el = styles_root.find(qn("w:latentStyles"))
	if latent_styles_el is not None:
		for latent_style_el in latent_styles_el.findall(qn("w:lsdException")):
			try:
				name = str(latent_style_el.get(qn("w:name")) or "").strip()
				if not name or _norm_style_key(name) in explicit_name_keys:
					continue
				items.append({
					"key": f"latent:{_norm_style_key(name)}",
					"name": name,
					"style_id": "(latent)",
					"type": "Word Default",
					"priority": _read_latent_style_ui_priority(latent_style_el),
				})
			except Exception:
				continue
	items.sort(key=lambda item: (int(item["priority"]), str(item["name"]).lower(), str(item["style_id"]).lower()))
	return items

def set_style_priorities_in_docx(input_path: Path, style_priority_by_key: dict[str, int]) -> int:
	doc = Document(str(input_path))
	styles_root = doc.part.styles._element
	changed = 0
	for style in doc.styles:
		try:
			style_id = str(getattr(style, "style_id", "") or "").strip()
			style_key = f"style:{_norm_style_key(style_id)}"
			if not style_id or style_key not in style_priority_by_key:
				continue
			new_prio = max(1, int(style_priority_by_key[style_key]))
			style_el = style._element
			ui_node = style_el.find(qn("w:uiPriority"))
			if ui_node is None:
				ui_node = OxmlElement("w:uiPriority")
				style_el.append(ui_node)
			old_val = str(ui_node.get(qn("w:val")) or "").strip()
			if old_val != str(new_prio):
				changed += 1
			ui_node.set(qn("w:val"), str(new_prio))
		except Exception:
			continue
	latent_styles_el = styles_root.find(qn("w:latentStyles"))
	if latent_styles_el is not None:
		for latent_style_el in latent_styles_el.findall(qn("w:lsdException")):
			try:
				name = str(latent_style_el.get(qn("w:name")) or "").strip()
				if not name:
					continue
				style_key = f"latent:{_norm_style_key(name)}"
				if style_key not in style_priority_by_key:
					continue
				new_prio = max(1, int(style_priority_by_key[style_key]))
				old_val = str(latent_style_el.get(qn("w:uiPriority")) or "").strip()
				if old_val != str(new_prio):
					changed += 1
				latent_style_el.set(qn("w:uiPriority"), str(new_prio))
			except Exception:
				continue
	doc.save(str(input_path))
	return changed

def group_consecutive_paragraphs_by_style(paragraphs: list, paragraph_indices: list[int] | None = None) -> list[list]:
	groups: list[list] = []
	current: list = []
	current_style = None
	current_index: int | None = None
	indices = paragraph_indices if paragraph_indices is not None else list(range(len(paragraphs)))

	for paragraph, paragraph_index in zip(paragraphs, indices):
		style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
		if not current:
			current = [paragraph]
			current_style = style_name
			current_index = paragraph_index
			continue

		if style_name == current_style and current_index is not None and paragraph_index == current_index + 1:
			current.append(paragraph)
			current_index = paragraph_index
		else:
			groups.append(current)
			current = [paragraph]
			current_style = style_name
			current_index = paragraph_index

	if current:
		groups.append(current)

	return groups

def extract_translated_paragraph_blocks(translated_text: str, paragraph_count: int) -> list[str]:
	blocks: list[str] = []
	for idx in range(1, paragraph_count + 1):
		token = f"PARA_{idx:04d}"
		pattern = rf"\[\[{token}\]\](.*?)\[\[/{token}\]\]"
		match = re.search(pattern, translated_text, flags=re.DOTALL)
		if not match:
			raise ValueError("Model output has invalid paragraph-anchor mapping.")
		blocks.append(match.group(1))
	return blocks

def translate_docx(
	input_path: Path,
	output_path: Path,
	translator: OpenAIParagraphTranslator,
	selected_styles: set[str],
	selected_char_styles: set[str] | None,
	target_language: str,
	debug_mode: bool = False,
	progress_callback=None,
	cancel_event: threading.Event | None = None,
	pause_event: threading.Event | None = None,
	pause_callback: Callable[[bool, int, int], None] | None = None,
) -> tuple[int, int, int, int, bool, list[dict[str, str]]]:
	doc = Document(str(input_path))
	all_paragraphs = collect_all_paragraphs(doc)
	entries = find_translatable_paragraph_entries(all_paragraphs, selected_styles)
	paragraphs = [paragraph for _, paragraph in entries]
	paragraph_indices = [idx for idx, _ in entries]
	snapshot_path = build_last_source_snapshot_path(input_path, target_language)

	def on_precheck_progress(processed: int, total: int, label: str) -> None:
		if progress_callback:
			progress_callback(
				processed,
				total,
				0,
				0,
				None,
				"Precheck",
				0,
				True,
				None,
			)

	incremental_result = apply_incremental_translation_seed(
		current_doc=doc,
		output_path=output_path,
		snapshot_path=snapshot_path,
		selected_styles=selected_styles,
		selected_char_styles=(selected_char_styles or set()),
		precheck_progress_callback=on_precheck_progress,
	)
	output_doc: Document | None = None
	output_all_paragraphs: list = []
	if incremental_result is not None:
		paragraphs, paragraph_indices, output_doc, output_all_paragraphs = incremental_result
	# Always use absolute paragraph indices so non-selected styles still split blocks.
	paragraph_groups = group_consecutive_paragraphs_by_style(paragraphs, paragraph_indices=paragraph_indices)
	total_candidates = len(paragraphs)
	total_blocks = len(paragraph_groups)
	processed_blocks = 0
	translated_count = 0
	processed_count = 0
	failed_count = 0
	cancelled = False
	failed_items: list[dict[str, str]] = []
	warning_count = 0
	warning_items: list[dict[str, str]] = []

	for group in paragraph_groups:
		if pause_event and pause_event.is_set():
			# Save a checkpoint so the current output file can be opened while paused.
			try:
				apply_document_language(doc, target_language)
				doc.save(str(output_path))
			except Exception:
				pass
			if pause_callback:
				pause_callback(True, processed_blocks, total_blocks)
			while pause_event.is_set():
				if cancel_event and cancel_event.is_set():
					cancelled = True
					break
				time.sleep(0.1)
			if pause_callback:
				pause_callback(False, processed_blocks, total_blocks)
			if cancelled:
				break

		if cancel_event and cancel_event.is_set():
			cancelled = True
			break

		group_infos: list[dict[str, object]] = []
		group_parts: list[str] = []
		group_hard_tokens: list[str] = []
		group_color_tokens: list[str] = []
		group_char_style_tokens: list[str] = []
		group_hard_elements: dict[str, dict[str, object]] = {}
		group_color_styles: dict[str, object] = {}
		group_char_style_map: dict[str, object] = {}
		counters = {"MATH": 0, "OBJ": 0, "CLR": 0, "CHST": 0}

		for idx, paragraph in enumerate(group, start=1):
			label = paragraph_label(paragraph)
			tagged_paragraph, hard_tokens, hard_elements, color_tokens, color_styles, char_style_tokens, char_style_map, source_debug = build_math_tagged_paragraph(
				paragraph,
				selected_char_styles=selected_char_styles,
			)

			mapped_hard_tokens: list[str] = []
			mapped_hard_elements: dict[str, dict[str, object]] = {}
			hard_replacements: dict[str, str] = {}
			for old_token in hard_tokens:
				kind = old_token[2: old_token.find("_")]
				counters[kind] = counters.get(kind, 0) + 1
				new_token = f"[[{kind}_{counters[kind]:04d}]]"
				hard_replacements[old_token] = new_token
				mapped_hard_tokens.append(new_token)
				mapped_hard_elements[new_token] = hard_elements[old_token]
				group_hard_tokens.append(new_token)
				group_hard_elements[new_token] = hard_elements[old_token]
			tagged_paragraph = replace_many_exact(tagged_paragraph, hard_replacements)

			mapped_color_tokens: list[str] = []
			mapped_color_styles: dict[str, object] = {}
			color_replacements: dict[str, str] = {}
			for old_token in color_tokens:
				counters["CLR"] = counters.get("CLR", 0) + 1
				new_token = f"CLR_{counters['CLR']:04d}"
				color_replacements[f"[[{old_token}]]"] = f"[[{new_token}]]"
				color_replacements[f"[[/{old_token}]]"] = f"[[/{new_token}]]"
				mapped_color_tokens.append(new_token)
				mapped_color_styles[new_token] = color_styles[old_token]
				group_color_tokens.append(new_token)
				group_color_styles[new_token] = color_styles[old_token]
			tagged_paragraph = replace_many_exact(tagged_paragraph, color_replacements)

			mapped_char_style_tokens: list[str] = []
			mapped_char_style_map: dict[str, object] = {}
			char_replacements: dict[str, str] = {}
			for old_token in char_style_tokens:
				counters["CHST"] = counters.get("CHST", 0) + 1
				new_token = f"CHST_{counters['CHST']:04d}"
				char_replacements[f"[[{old_token}]]"] = f"[[{new_token}]]"
				char_replacements[f"[[/{old_token}]]"] = f"[[/{new_token}]]"
				mapped_char_style_tokens.append(new_token)
				mapped_char_style_map[new_token] = char_style_map[old_token]
				group_char_style_tokens.append(new_token)
				group_char_style_map[new_token] = char_style_map[old_token]
			tagged_paragraph = replace_many_exact(tagged_paragraph, char_replacements)

			para_token = f"PARA_{idx:04d}"
			group_parts.append(f"[[{para_token}]]{tagged_paragraph}[[/{para_token}]]")
			group_infos.append(
				{
					"paragraph": paragraph,
					"label": label,
					"tagged_paragraph": tagged_paragraph,
					"hard_tokens": mapped_hard_tokens,
					"hard_elements": mapped_hard_elements,
					"color_tokens": mapped_color_tokens,
					"color_styles": mapped_color_styles,
					"char_style_tokens": mapped_char_style_tokens,
					"char_style_map": mapped_char_style_map,
					"source_debug": source_debug,
				}
			)

		block_size = len(group_infos)
		if block_size == 1:
			block_tagged_paragraph = str(group_infos[0]["tagged_paragraph"])
		else:
			block_tagged_paragraph = "".join(group_parts)
		block_label = f"Block {block_size} Para. | {group_infos[0]['label']}"
		current_attempt = 0

		def on_attempt(attempt: int) -> None:
			nonlocal current_attempt
			current_attempt = attempt
			if progress_callback:
				progress_callback(
					processed_blocks,
					total_blocks,
					translated_count,
					failed_count,
					None,
					block_label,
					attempt,
					True,
					None,
				)

		try:
			translated_block = translator.translate_tagged_paragraph(
				tagged_paragraph=block_tagged_paragraph,
				hard_tokens=group_hard_tokens,
				color_tokens=group_color_tokens,
				char_style_tokens=group_char_style_tokens,
				char_style_map=group_char_style_map,
				attempt_callback=on_attempt,
				post_validate=(
					(lambda text: extract_translated_paragraph_blocks(text, len(group_infos)))
					if len(group_infos) > 1
					else None
				),
			)
		except Exception as exc:
			block_progress = min(processed_blocks + 1, total_blocks) if total_blocks else 0
			msg = str(exc).lower()

			# Single-paragraph hard-anchor fallback: positional placement, shown as warning
			if isinstance(exc, HardAnchorFallbackWarning) and len(group_infos) == 1:
				info = group_infos[0]
				apply_translated_with_appended_hard_elements(
					paragraph=info["paragraph"],
					translated_plain=exc.fallback_text,
					hard_elements=info["hard_elements"],
					color_styles=info["color_styles"],
					char_style_map=info.get("char_style_map"),
				)
				translated_count += 1
				processed_count += 1
				warning_count += 1
				warning_items.append({"paragraph": str(info["label"]), "error": str(exc), "attempts": str(current_attempt)})
				if progress_callback:
					progress_callback(block_progress, total_blocks, translated_count, failed_count, exc, str(info["label"]), current_attempt, False, None)
				processed_blocks = block_progress
				continue

			can_split_fallback = (len(group_infos) > 1) and ("color-anchor" in msg or "anchor" in msg)

			if can_split_fallback:
				for info in group_infos:
					label = str(info["label"])
					paragraph = info["paragraph"]
					tagged_single = str(info["tagged_paragraph"])
					hard_tokens = info["hard_tokens"]
					hard_elements = info["hard_elements"]
					color_tokens = info["color_tokens"]
					color_styles = info["color_styles"]
					char_style_tokens_single = list(info.get("char_style_tokens", []))
					char_style_map_single = dict(info.get("char_style_map", {}))
					source_debug = info["source_debug"]

					single_attempt = 0
					def on_attempt_single(attempt: int) -> None:
						nonlocal single_attempt
						single_attempt = attempt
						if progress_callback:
							progress_callback(
								processed_blocks,
								total_blocks,
								translated_count,
								failed_count,
								None,
								f"Fallback | {label}",
								attempt,
								True,
								None,
							)

					try:
						single_translated = translator.translate_tagged_paragraph(
							tagged_paragraph=tagged_single,
							hard_tokens=hard_tokens,
							color_tokens=color_tokens,
							char_style_tokens=char_style_tokens_single,
							char_style_map=char_style_map_single,
							attempt_callback=on_attempt_single,
						)
					except HardAnchorFallbackWarning as warn_single:
						apply_translated_with_appended_hard_elements(
							paragraph=paragraph,
							translated_plain=warn_single.fallback_text,
							hard_elements=hard_elements,
							color_styles=color_styles,
							char_style_map=char_style_map_single,
						)
						translated_count += 1
						processed_count += 1
						warning_count += 1
						warning_items.append({"paragraph": label, "error": str(warn_single), "attempts": str(single_attempt)})
						if progress_callback:
							progress_callback(block_progress, total_blocks, translated_count, failed_count, warn_single, label, single_attempt, False, None)
						continue
					except Exception as exc_single:
						failed_count += 1
						processed_count += 1
						failed_items.append({"paragraph": label, "error": str(exc_single), "attempts": str(single_attempt)})
						if progress_callback:
							debug_payload = None
							if debug_mode:
								debug_payload = {
									"paragraph": label,
									"hard_token_count": len(hard_tokens),
									"color_token_count": len(color_tokens),
									"source_bold_runs": int(source_debug.get("source_bold_runs", 0)),
									"source_named_bold_runs": int(source_debug.get("source_named_bold_runs", 0)),
									"source_bold_segments": int(source_debug.get("source_bold_segments", 0)),
									"source_colored_segments": int(source_debug.get("source_colored_segments", 0)),
									"source_has_named_bold": bool(source_debug.get("source_has_named_bold", False)),
									"used_run_styles": list(source_debug.get("used_run_styles", [])),
									"used_char_styles": list(source_debug.get("used_char_styles", [])),
									"used_color_values": list(source_debug.get("used_color_values", [])),
									"error": str(exc_single),
								}
							progress_callback(
								block_progress,
								total_blocks,
								translated_count,
								failed_count,
								exc_single,
								label,
								single_attempt,
								False,
								debug_payload,
							)
						continue

					replace_debug = replace_paragraph_with_translated(
						paragraph=paragraph,
						translated_text=single_translated,
						hard_tokens=hard_tokens,
						hard_elements=hard_elements,
						color_styles=color_styles,
						char_style_map=char_style_map_single,
					)
					translated_count += 1
					processed_count += 1
					if progress_callback:
						debug_payload = None
						if debug_mode:
							debug_payload = {
								"paragraph": label,
								"hard_token_count": len(hard_tokens),
								"color_token_count": len(color_tokens),
								"source_bold_runs": int(source_debug.get("source_bold_runs", 0)),
								"source_named_bold_runs": int(source_debug.get("source_named_bold_runs", 0)),
								"source_bold_segments": int(source_debug.get("source_bold_segments", 0)),
								"source_colored_segments": int(source_debug.get("source_colored_segments", 0)),
								"source_has_named_bold": bool(source_debug.get("source_has_named_bold", False)),
								"used_run_styles": list(source_debug.get("used_run_styles", [])),
								"used_char_styles": list(source_debug.get("used_char_styles", [])),
								"used_color_values": list(source_debug.get("used_color_values", [])),
								"preferred_style": str(replace_debug.get("preferred_style", "-")),
								"style_fett_exists": bool(replace_debug.get("style_fett_exists", False)),
								"style_bold_exists": bool(replace_debug.get("style_bold_exists", False)),
								"output_bold_segments": int(replace_debug.get("output_bold_segments", 0)),
								"applied_by_style": int(replace_debug.get("applied_by_style", 0)),
								"applied_direct": int(replace_debug.get("applied_direct", 0)),
							}
						progress_callback(
							block_progress,
							total_blocks,
							translated_count,
							failed_count,
							None,
							label,
							single_attempt,
							False,
							debug_payload,
						)

				processed_blocks = block_progress
				continue

			for info in group_infos:
				failed_count += 1
				processed_count += 1
				label = str(info["label"])
				source_debug = info["source_debug"]
				color_tokens = info["color_tokens"]
				hard_tokens = info["hard_tokens"]
				failed_items.append({"paragraph": label, "error": str(exc), "attempts": str(current_attempt)})
				if progress_callback:
					debug_payload = None
					if debug_mode:
						debug_payload = {
							"paragraph": label,
							"hard_token_count": len(hard_tokens),
							"color_token_count": len(color_tokens),
							"source_bold_runs": int(source_debug.get("source_bold_runs", 0)),
							"source_named_bold_runs": int(source_debug.get("source_named_bold_runs", 0)),
							"source_bold_segments": int(source_debug.get("source_bold_segments", 0)),
							"source_colored_segments": int(source_debug.get("source_colored_segments", 0)),
							"source_has_named_bold": bool(source_debug.get("source_has_named_bold", False)),
							"used_run_styles": list(source_debug.get("used_run_styles", [])),
							"used_char_styles": list(source_debug.get("used_char_styles", [])),
							"used_color_values": list(source_debug.get("used_color_values", [])),
							"error": str(exc),
						}
					progress_callback(
						block_progress,
						total_blocks,
						translated_count,
						failed_count,
						exc,
						label,
						current_attempt,
						False,
						debug_payload,
					)
			processed_blocks = block_progress
			continue

		try:
			if len(group_infos) > 1:
				translated_blocks = extract_translated_paragraph_blocks(translated_block, len(group_infos))
			else:
				translated_blocks = [translated_block]
		except Exception as exc:
			block_progress = min(processed_blocks + 1, total_blocks) if total_blocks else 0
			for info in group_infos:
				failed_count += 1
				processed_count += 1
				label = str(info["label"])
				failed_items.append({"paragraph": label, "error": str(exc), "attempts": str(current_attempt)})
				if progress_callback:
					progress_callback(
						block_progress,
						total_blocks,
						translated_count,
						failed_count,
						exc,
						label,
						current_attempt,
						False,
						None,
					)
			processed_blocks = block_progress
			continue

		block_progress = min(processed_blocks + 1, total_blocks) if total_blocks else 0

		for info, translated_text in zip(group_infos, translated_blocks):
			paragraph = info["paragraph"]
			label = str(info["label"])
			hard_tokens = info["hard_tokens"]
			hard_elements = info["hard_elements"]
			color_styles = info["color_styles"]
			color_tokens = info["color_tokens"]
			source_debug = info["source_debug"]

			replace_debug = replace_paragraph_with_translated(
				paragraph=paragraph,
				translated_text=translated_text,
				hard_tokens=hard_tokens,
				hard_elements=hard_elements,
				color_styles=color_styles,
				char_style_map=info.get("char_style_map"),
			)

			translated_count += 1
			processed_count += 1
			if progress_callback:
				debug_payload = None
				if debug_mode:
					debug_payload = {
						"paragraph": label,
						"hard_token_count": len(hard_tokens),
						"color_token_count": len(color_tokens),
						"source_bold_runs": int(source_debug.get("source_bold_runs", 0)),
						"source_named_bold_runs": int(source_debug.get("source_named_bold_runs", 0)),
						"source_bold_segments": int(source_debug.get("source_bold_segments", 0)),
						"source_colored_segments": int(source_debug.get("source_colored_segments", 0)),
						"source_has_named_bold": bool(source_debug.get("source_has_named_bold", False)),
						"used_run_styles": list(source_debug.get("used_run_styles", [])),
						"used_char_styles": list(source_debug.get("used_char_styles", [])),
						"used_color_values": list(source_debug.get("used_color_values", [])),
						"preferred_style": str(replace_debug.get("preferred_style", "-")),
						"style_fett_exists": bool(replace_debug.get("style_fett_exists", False)),
						"style_bold_exists": bool(replace_debug.get("style_bold_exists", False)),
						"output_bold_segments": int(replace_debug.get("output_bold_segments", 0)),
						"applied_by_style": int(replace_debug.get("applied_by_style", 0)),
						"applied_direct": int(replace_debug.get("applied_direct", 0)),
					}
				progress_callback(
					block_progress,
					total_blocks,
					translated_count,
					failed_count,
					None,
					label,
					current_attempt,
					False,
					debug_payload,
				)

		processed_blocks = block_progress

	if output_doc is not None:
		# Incremental mode: write translated paragraphs into the existing output file.
		# Unchanged blocks are not touched at all in output_doc.
		for idx in paragraph_indices:
			if idx < len(output_all_paragraphs):
				replace_paragraph_xml(output_all_paragraphs[idx], all_paragraphs[idx])
		apply_document_language(output_doc, target_language)
		output_doc.save(str(output_path))
	else:
		apply_document_language(doc, target_language)
		doc.save(str(output_path))
	snapshot_save_error: str | None = None
	if not cancelled:
		try:
			shutil.copy2(str(input_path), str(snapshot_path))
		except Exception as exc:
			snapshot_save_error = str(exc)
			warning_count += 1
			failed_items.append({
				"paragraph": "Snapshot",
				"error": f"Snapshot could not be saved: {snapshot_path}",
				"attempts": "-",
			})
	if snapshot_save_error:
		failed_items.append({
			"paragraph": "Snapshot",
			"error": f"Snapshot save exception: {snapshot_save_error}",
			"attempts": "-",
		})
	return total_blocks, processed_blocks, failed_count, warning_count, cancelled, failed_items

class TranslatorApp:
	def __init__(self, root: tk.Tk) -> None:
		self.root = root
		self.root.title("DOCX Translator")
		self.root.geometry("1320x760")
		self.root.minsize(1220, 720)
		self.root.configure(bg="#eef2f7")

		self.queue: Queue = Queue()
		self.worker: threading.Thread | None = None
		self.model_loader: threading.Thread | None = None
		self.model_price_loader: threading.Thread | None = None
		self.cancel_event = threading.Event()
		self.pause_event = threading.Event()
		self.is_running = False
		self.is_paused = False
		self._taskbar = None
		self.models_loading = False
		self.prices_loading = False
		self.error_log_text: tk.Text | None = None
		self.error_count = 0
		self.last_output: Path | None = None
		self.last_output_exists = False
		self.model_choices: list[str] = []
		self.model_choice_labels: list[str] = []
		self.model_label_to_id: dict[str, str] = {}
		self.model_id_to_label: dict[str, str] = {}
		self.model_picker_window: tk.Toplevel | None = None
		self.model_picker_tree: ttk.Treeview | None = None
		self.model_picker_sort_column = "model"
		self.model_picker_sort_desc = False
		self.settings = load_gui_settings()
		self.language_codes = _clean_language_codes(self.settings.get("language_codes"))
		saved_styles = self.settings.get("selected_styles")
		saved_char_styles = self.settings.get("selected_char_styles")
		self.selected_styles: set[str] = set(saved_styles) if saved_styles else set(TARGET_STYLE_NAMES)
		self.selected_char_styles: set[str] = set(saved_char_styles) if saved_char_styles else set()
		self.glossary_pairs: dict[str, list[dict[str, str]]] = normalize_glossary_pairs(self.settings.get("glossary_pairs"))
		self.available_styles: list[str] = []
		self.available_char_styles: list[str] = []
		self.styles_scanning = False

		self.file_var = tk.StringVar(value=str(self.settings.get("input_file", "")))
		default_model = str(self.settings.get("model") or os.getenv("OPENAI_MODEL") or "")
		default_source = _canon_lang(str(self.settings.get("source_language") or "de"))
		default_target = _canon_lang(str(self.settings.get("target_language") or "en-US"))
		if default_source not in self.language_codes:
			default_source = self.language_codes[0]
		if default_target not in self.language_codes:
			default_target = self.language_codes[1] if len(self.language_codes) > 1 else self.language_codes[0]
		saved_api_key = str(self.settings.get("api_key") or os.getenv("OPENAI_API_KEY", ""))
		self.model_var = tk.StringVar(value=default_model)
		self.model_display_var = tk.StringVar(value=(build_model_choice_label(default_model) if default_model else ""))
		self.api_key_var = tk.StringVar(value=saved_api_key)
		self.source_lang_var = tk.StringVar(value=default_source)
		self.target_lang_var = tk.StringVar(value=default_target)
		self.theme_var = tk.StringVar(value=_normalize_theme_mode(self.settings.get("theme_mode")))
		self.debug_var = tk.BooleanVar(value=bool(self.settings.get("debug_mode", False)))
		self.styles_info_var = tk.StringVar(value="Select a file first")
		self.char_styles_info_var = tk.StringVar(value="Select a file first")
		self.glossary_info_var = tk.StringVar(value="Glossary: 0 entries")
		self.status_var = tk.StringVar(value="Select a DOCX file and a model.")
		self.progress_var = tk.DoubleVar(value=0)
		self.warning_counter_var = tk.StringVar(value="Warnings: 0")
		self.error_counter_var = tk.StringVar(value="Errors: 0")
		self.token_counter_var = tk.StringVar(value="Tokens: In: 0 | Out: 0")
		self.cost_counter_var = tk.StringVar(value="Cost: n/a")
		self.live_info_var = tk.StringVar(value="Live-Info: -")
		self.api_request_info_var = tk.StringVar(value="API Requests: 0 | Req-ID: - | Resp-ID: -")
		self.warning_count_live = 0
		self.error_count_live = 0
		self.usage_prompt_tokens_live = 0
		self.usage_completion_tokens_live = 0
		self.usage_total_tokens_live = 0
		self.usage_cost_usd_live = 0.0
		self.usage_cost_eur_live = 0.0
		self.usage_eur_rate_live: float | None = None
		self.usage_pricing_known_live = False
		self.api_request_count_live = 0
		self.last_request_id_live = "-"
		self.last_response_id_live = "-"
		self.current_block_label = ""
		self.current_block_start_tokens = 0
		self.current_block_start_cost_usd = 0.0
		self.current_block_start_cost_eur = 0.0

		self._configure_style()
		self._build_ui()
		self._init_taskbar()
		self._set_model_choices(self.model_choices)
		saved_file = self.file_var.get().strip()
		if saved_file and Path(saved_file).exists():
			self._scan_styles_from_file(Path(saved_file), on_startup=True)
		else:
			self._update_styles_info()
			self._update_char_styles_info()
		self._update_glossary_info()
		self._request_model_list(on_startup=True)
		self.root.after(100, self._poll_queue)

	def _configure_style(self) -> None:
		self._apply_theme()
		style = ttk.Style()

		# Font definitions for custom labels
		style.configure("Header.TLabel", font=("Segoe UI Semibold", 20))
		style.configure("Sub.TLabel", font=("Segoe UI", 10))
		style.configure("Label.TLabel", font=("Segoe UI Semibold", 10))
		style.configure("Status.TLabel", font=("Segoe UI", 10))
		style.configure("StatusMono.TLabel", font=("Cascadia Code", 10))
		style.configure("ErrorTitle.TLabel", font=("Segoe UI Semibold", 11))
		style.configure("WarningCounter.TLabel", font=("Segoe UI Semibold", 10))
		style.configure("ErrorCounter.TLabel", font=("Segoe UI Semibold", 10))
		style.configure("UsageCounter.TLabel", font=("Segoe UI Semibold", 10))
		style.configure("LiveInfo.TLabel", font=("Cascadia Code", 9))

		# Custom button styles to enhance Sun-Valley theme
		style.configure("Primary.TButton", font=("Segoe UI Semibold", 10), padding=(14, 8))
		style.configure("Secondary.TButton", font=("Segoe UI", 10), padding=(10, 7))

		style.configure(
			"Modern.Horizontal.TProgressbar",
			thickness=14,
		)

	def _apply_theme(self) -> None:
		theme_mode = _normalize_theme_mode(self.theme_var.get() if hasattr(self, "theme_var") else DEFAULT_THEME_MODE)
		if hasattr(self, "theme_var"):
			self.theme_var.set(theme_mode)
		svc_ttk.set_theme(_resolve_sv_ttk_theme(theme_mode))

	def _build_ui(self) -> None:
		outer = ttk.Frame(self.root, padding=GS)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="DOCX Translation Studio", style="Header.TLabel").pack(anchor="w")
		ttk.Label(
			outer,
			text="Translates paragraph by paragraph, preserving formatting and inline objects.",
			style="Sub.TLabel",
		).pack(anchor="w")

		body = ttk.PanedWindow(outer, orient="horizontal")
		body.pack(fill="both", expand=True)

		# ── LEFT PANEL ────────────────────────────────────────────────────────
		left = ttk.Frame(body, style="Card.TFrame", padding=GS)
		left.columnconfigure(0, weight=1)
		left.columnconfigure(1, weight=0)
		body.add(left, weight=1)

		# ── RIGHT PANEL (expands) ───────────────────────────────────────────────
		right = ttk.Frame(body, style="Card.TFrame", padding=GS)
		right.columnconfigure(0, weight=1)
		body.add(right, weight=1)

		# ── Left: File ─────────────────────────────────────────────────────────────────
		r = 0
		ttk.Label(left, text="Input DOCX", style="Label.TLabel").grid(row=r, column=0, sticky="w")
		r += 1
		file_entry = ttk.Entry(left, textvariable=self.file_var, font=("Segoe UI", 10))
		file_entry.grid(row=r, column=0, sticky="ew")
		self.file_button = ttk.Button(
			left, text="Choose File", style="Secondary.TButton", width=17, command=self._choose_file
		)
		self.file_button.grid(row=r, column=1, sticky="ew")
		r += 1
		ttk.Separator(left, orient="horizontal").grid(row=r, column=0, columnspan=2, sticky="ew")
		r += 1

		# ── Left: Styles ───────────────────────────────────────────────────────────────────
		ttk.Label(left, textvariable=self.styles_info_var, style="Status.TLabel").grid(row=r, column=0, sticky="w")
		self.styles_button = ttk.Button(
			left, text="Styles", style="Secondary.TButton", width=17,
			command=self._open_style_selector
		)
		self.styles_button.grid(row=r, column=1, sticky="ew")
		self.styles_button.configure(state="disabled")
		r += 1
		ttk.Label(left, textvariable=self.char_styles_info_var, style="Status.TLabel").grid(row=r, column=0, sticky="w")
		self.char_styles_button = ttk.Button(
			left, text="Char Styles", style="Secondary.TButton", width=17,
			command=self._open_char_style_selector
		)
		self.char_styles_button.grid(row=r, column=1, sticky="ew")
		self.char_styles_button.configure(state="disabled")
		r += 1
		self.cleaner_button = ttk.Button(
			left, text="Clean Up", style="Secondary.TButton", width=17,
			command=self._open_style_cleaner
		)
		self.cleaner_button.grid(row=r, column=1, sticky="ew")
		self.cleaner_button.configure(state="disabled")
		r += 1
		self.sort_styles_button = ttk.Button(
			left, text="Sort Styles", style="Secondary.TButton", width=17,
			command=self._open_style_priority_editor
		)
		self.sort_styles_button.grid(row=r, column=1, sticky="ew")
		self.sort_styles_button.configure(state="disabled")
		r += 1
		ttk.Separator(left, orient="horizontal").grid(row=r, column=0, columnspan=2, sticky="ew")
		r += 1

		# ── Left: Languages ──────────────────────────────────────────────────────────────────
		lang_frame = ttk.Frame(left, padding=GS)
		lang_frame.grid(row=r, column=0, columnspan=2, sticky="ew")
		lang_frame.columnconfigure(0, weight=1)
		lang_frame.columnconfigure(1, weight=1)
		ttk.Label(lang_frame, text="Source Language", style="Label.TLabel").grid(row=0, column=0, sticky="w")
		ttk.Label(lang_frame, text="Target Language", style="Label.TLabel").grid(row=0, column=1, sticky="w")
		self.source_combo = ttk.Combobox(
			lang_frame, textvariable=self.source_lang_var, values=self.language_codes,
			state="readonly", font=("Segoe UI", 10)
		)
		self.source_combo.grid(row=1, column=0, sticky="ew")
		self.source_combo.bind("<<ComboboxSelected>>", self._on_language_change)
		self.target_combo = ttk.Combobox(
			lang_frame, textvariable=self.target_lang_var, values=self.language_codes,
			state="readonly", font=("Segoe UI", 10)
		)
		self.target_combo.grid(row=1, column=1, sticky="ew")
		self.target_combo.bind("<<ComboboxSelected>>", self._on_language_change)
		self.source_combo.configure(values=self.language_codes)
		self.target_combo.configure(values=self.language_codes)
		self.languages_button = ttk.Button(
			lang_frame,
			text="Manage Languages",
			style="Secondary.TButton",
			command=self._open_language_manager,
		)
		self.languages_button.grid(row=2, column=0, columnspan=2, sticky="ew")
		r += 1

		# ── Left: Glossary ──────────────────────────────────────────────────────────────────
		ttk.Label(left, textvariable=self.glossary_info_var, style="Status.TLabel").grid(row=r, column=0, sticky="w")
		self.glossary_button = ttk.Button(
			left, text="Edit Glossary", style="Secondary.TButton", width=17,
			command=self._open_glossary_editor
		)
		self.glossary_button.grid(row=r, column=1, sticky="ew")
		r += 1
		ttk.Separator(left, orient="horizontal").grid(row=r, column=0, columnspan=2, sticky="ew")
		r += 1

		# ── Left: Model ─────────────────────────────────────────────────────────────────────
		ttk.Label(left, text="OpenAI Model", style="Label.TLabel").grid(row=r, column=0, columnspan=2, sticky="w")
		r += 1
		self.model_selector_frame = ttk.Frame(left, padding=GS)
		self.model_selector_frame.grid(row=r, column=0, columnspan=2, sticky="ew")
		self.model_selector_frame.columnconfigure(0, weight=1)
		self.model_selector_frame.bind("<Configure>", self._on_model_selector_resize)
		self.model_display_table = ttk.Treeview(
			self.model_selector_frame,
			columns=("model", "input", "cached", "output"),
			show="headings",
			height=1,
			selectmode="none",
		)
		self.model_display_table.heading("model", text="Model")
		self.model_display_table.heading("input", text="Input")
		self.model_display_table.heading("cached", text="Cached")
		self.model_display_table.heading("output", text="Output")
		self.model_display_table.column("model", anchor="w", width=260, stretch=True)
		self.model_display_table.column("input", anchor="e", width=120, stretch=True)
		self.model_display_table.column("cached", anchor="e", width=120, stretch=True)
		self.model_display_table.column("output", anchor="e", width=120, stretch=True)
		self.model_display_table.grid(row=0, column=0, sticky="ew")
		self.model_display_table.bind("<Button-1>", self._open_model_picker_from_event)
		self.model_picker_button = ttk.Button(
			self.model_selector_frame,
			text="v",
			style="Secondary.TButton",
			width=4,
			command=self._open_model_picker,
		)
		self.model_picker_button.grid(row=0, column=1, sticky="ns")
		self._resize_model_tree_columns(self.model_display_table, self.model_selector_frame.winfo_width())
		r += 1
		model_actions = ttk.Frame(left, padding=GS)
		model_actions.grid(row=r, column=0, columnspan=2, sticky="ew")
		model_actions.columnconfigure(0, weight=1)
		model_actions.columnconfigure(1, weight=1)
		model_actions.columnconfigure(2, weight=1)
		self.refresh_models_button = ttk.Button(
			model_actions, text="Load Models", style="Secondary.TButton",
			command=self._request_model_list
		)
		self.refresh_models_button.grid(row=0, column=0, sticky="ew")
		self.refresh_prices_button = ttk.Button(
			model_actions, text="Update Model Prices", style="Secondary.TButton",
			command=self._request_model_price_update
		)
		self.refresh_prices_button.grid(row=0, column=1, sticky="ew")
		self.settings_button = ttk.Button(
			model_actions, text="Settings", style="Secondary.TButton",
			command=self._open_settings_window
		)
		self.settings_button.grid(row=0, column=2, sticky="ew")
		r += 1
		ttk.Separator(left, orient="horizontal").grid(row=r, column=0, columnspan=2, sticky="ew")
		r += 1

		# ── Left: Start / Pause / Cancel ───────────────────────────────────────────────────
		button_row = ttk.Frame(left, padding=GS)
		button_row.grid(row=r, column=0, columnspan=2, sticky="ew")
		button_row.columnconfigure(0, weight=1)
		button_row.columnconfigure(1, weight=1)
		button_row.columnconfigure(2, weight=1)
		self.start_button = ttk.Button(
			button_row, text="Start Translation", command=self._start
		)
		self.start_button.grid(row=0, column=0, sticky="ew")
		self.pause_button = ttk.Button(
			button_row, text="Pause", command=self._toggle_pause
		)
		self.pause_button.grid(row=0, column=1, sticky="ew")
		self.pause_button.configure(state="disabled")
		self.cancel_button = ttk.Button(
			button_row, text="Cancel", command=self._cancel
		)
		self.cancel_button.grid(row=0, column=2, sticky="ew")
		self.cancel_button.configure(state="disabled")
		self.debug_check = ttk.Checkbutton(
			button_row, text="Debug Mode (Anchor Diagnostics)", variable=self.debug_var,
			command=self._persist_model_selection
		)
		self.debug_check.grid(row=1, column=0, columnspan=3, sticky="w")
		r += 1
		left.rowconfigure(r, weight=1)

		# ── Right: Progress ──────────────────────────────────────────────────────────────
		rr = 0
		ttk.Label(right, text="Progress", style="Label.TLabel").grid(row=rr, column=0, sticky="w")
		rr += 1
		self.progress = ttk.Progressbar(
			right, style="Modern.Horizontal.TProgressbar", orient="horizontal",
			mode="determinate", maximum=100, variable=self.progress_var
		)
		self.progress.grid(row=rr, column=0, sticky="ew")
		rr += 1

		# ── Right: Counters ───────────────────────────────────────────────────────────────
		counter_outer = ttk.Frame(right, padding=GS)
		counter_outer.grid(row=rr, column=0, sticky="ew")
		counter_outer.columnconfigure(0, weight=1)
		counter_outer.columnconfigure(1, weight=1)
		self.warning_counter_label = ttk.Label(
			counter_outer, textvariable=self.warning_counter_var, style="WarningCounter.TLabel"
		)
		self.warning_counter_label.grid(row=0, column=0, sticky="w")
		self.error_counter_label = ttk.Label(
			counter_outer, textvariable=self.error_counter_var, style="ErrorCounter.TLabel"
		)
		self.error_counter_label.grid(row=1, column=0, sticky="w")
		self.token_counter_label = ttk.Label(
			counter_outer, textvariable=self.token_counter_var, style="UsageCounter.TLabel"
		)
		self.token_counter_label.grid(row=0, column=1, sticky="w")
		self.cost_counter_label = ttk.Label(
			counter_outer, textvariable=self.cost_counter_var, style="UsageCounter.TLabel"
		)
		self.cost_counter_label.grid(row=1, column=1, sticky="w")
		rr += 1

		# ── Right: Status / Live info ─────────────────────────────────────────────────
		ttk.Separator(right, orient="horizontal").grid(row=rr, column=0, sticky="ew")
		rr += 1
		self.status_label = ttk.Label(right, textvariable=self.status_var, style="StatusMono.TLabel")
		self.status_label.grid(row=rr, column=0, sticky="ew")
		rr += 1
		self.live_info_label = ttk.Label(right, textvariable=self.live_info_var, style="LiveInfo.TLabel")
		self.live_info_label.grid(row=rr, column=0, sticky="ew")
		rr += 1
		self.api_request_info_label = ttk.Label(
			right, textvariable=self.api_request_info_var, style="LiveInfo.TLabel"
		)
		self.api_request_info_label.grid(row=rr, column=0, sticky="ew")
		rr += 1

		# ── Right: Reveal button ──────────────────────────────────────────────────────────
		self.reveal_button = ttk.Button(
			right, text="Open in Explorer", style="Secondary.TButton",
			command=self._open_output_folder
		)
		self.reveal_button.grid(row=rr, column=0, sticky="ew")
		self.reveal_button.grid_remove()
		rr += 1

		# ── Right: Error log (expands) ────────────────────────────────────────────────────
		right.rowconfigure(rr, weight=1)
		error_row = ttk.Frame(right, padding=GS)
		error_row.grid(row=rr, column=0, sticky="nsew")
		error_row.columnconfigure(0, weight=1)
		error_row.rowconfigure(2, weight=1)
		ttk.Label(error_row, text="Error Log", style="ErrorTitle.TLabel").grid(row=0, column=0, sticky="w")
		controls_row = ttk.Frame(error_row, padding=GS)
		controls_row.grid(row=1, column=0, sticky="ew")
		self.clear_log_button = ttk.Button(
			controls_row, text="Clear Log", style="Secondary.TButton", command=self._clear_error_log
		)
		self.clear_log_button.pack(side="right")
		log_wrap = ttk.Frame(error_row, padding=GS)
		log_wrap.grid(row=2, column=0, sticky="nsew")
		log_wrap.rowconfigure(0, weight=1)
		log_wrap.columnconfigure(0, weight=1)
		self.error_log_text = tk.Text(
			log_wrap,
			wrap="word",
			font=("Cascadia Code", 10),
			bg="#111827",
			fg="#e5e7eb",
			insertbackground="#e5e7eb",
			relief="flat",
			padx=10,
			pady=10,
		)
		self.error_log_text.grid(row=0, column=0, sticky="nsew")
		scroll = ttk.Scrollbar(log_wrap, orient="vertical", command=self.error_log_text.yview)
		scroll.grid(row=0, column=1, sticky="ns")
		self.error_log_text.configure(yscrollcommand=scroll.set)
		self.error_log_text.tag_configure("entry_header", foreground="#fca5a5", font=("Cascadia Code", 10, "bold"))
		self.error_log_text.tag_configure("warn_header", foreground="#facc15", font=("Cascadia Code", 10, "bold"))
		self.error_log_text.tag_configure("info_header", foreground="#86efac", font=("Cascadia Code", 10, "bold"))
		self.error_log_text.tag_configure("entry_meta", foreground="#93c5fd")
		self.error_log_text.tag_configure("warn_meta", foreground="#fde68a")
		self.error_log_text.tag_configure("info_meta", foreground="#bbf7d0")
		self.error_log_text.tag_configure("entry_body", foreground="#e5e7eb")
		self.error_log_text.tag_configure("entry_sep", foreground="#6b7280")
		self.error_log_text.configure(state="disabled")

		for widget in left.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in right.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in counter_outer.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in button_row.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in lang_frame.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in model_actions.winfo_children():
			widget.grid(padx=GS, pady=GS)
		for widget in self.model_selector_frame.winfo_children():
			widget.grid(padx=GS, pady=GS)

	def _save_settings(self, *, input_file: str | None = None, api_key: str | None = None, language_codes: list[str] | None = None, theme_mode: str | None = None) -> None:
		save_gui_settings(
			input_file if input_file is not None else self.file_var.get().strip(),
			self.model_var.get().strip(),
			list(self.selected_styles),
			list(self.selected_char_styles),
			self.source_lang_var.get().strip(),
			self.target_lang_var.get().strip(),
			theme_mode if theme_mode is not None else self.theme_var.get().strip(),
			self.debug_var.get(),
			self.glossary_pairs,
			api_key if api_key is not None else self.api_key_var.get().strip(),
			language_codes=language_codes,
		)

	def _choose_file(self) -> None:
		file_path = filedialog.askopenfilename(
			title="Select DOCX File",
			filetypes=[("Word Document", "*.docx")],
		)
		if file_path:
			self.file_var.set(file_path)
			self.available_styles = []
			self.available_char_styles = []
			self.selected_styles = set()
			self.selected_char_styles = set()
			self._update_styles_info()
			self._update_char_styles_info()
			self._sync_controls()
			self._save_settings(input_file=file_path)
			self._scan_styles_from_file(Path(file_path))
			self.status_var.set("File selected. Reading styles...")

	def _sync_controls(self) -> None:
		if self.is_running:
			self.start_button.configure(state="disabled")
			self.cancel_button.configure(state="normal")
			self.pause_button.configure(state="normal")
			self.pause_button.configure(text=("Resume" if self.is_paused else "Pause"))
			self.file_button.configure(state="disabled")
			self.source_combo.configure(state="disabled")
			self.target_combo.configure(state="disabled")
			self.languages_button.configure(state="disabled")
			self.model_picker_button.configure(state="disabled")
			self._close_model_picker()
			self.settings_button.configure(state="disabled")
			self.refresh_models_button.configure(state="disabled")
			self.refresh_prices_button.configure(state="disabled")
			self.styles_button.configure(state="disabled")
			self.char_styles_button.configure(state="disabled")
			self.cleaner_button.configure(state="disabled")
			self.sort_styles_button.configure(state="disabled")
			self.glossary_button.configure(state="disabled")
			return

		self.cancel_button.configure(state="disabled")
		self.pause_button.configure(state="disabled", text="Pause")
		self.file_button.configure(state="normal")
		self.styles_button.configure(
			state="normal" if (self.available_styles and not self.styles_scanning) else "disabled"
		)
		self.char_styles_button.configure(
			state="normal" if (self.available_char_styles and not self.styles_scanning) else "disabled"
		)
		self.cleaner_button.configure(
			state="normal" if (self.file_var.get().strip() and not self.styles_scanning) else "disabled"
		)
		self.sort_styles_button.configure(
			state="normal" if (self.file_var.get().strip() and not self.styles_scanning) else "disabled"
		)
		self.glossary_button.configure(state="normal")

		if self.models_loading:
			self.start_button.configure(state="disabled")
			self.source_combo.configure(state="disabled")
			self.target_combo.configure(state="disabled")
			self.languages_button.configure(state="disabled")
			self.model_picker_button.configure(state="disabled")
			self._close_model_picker()
			self.char_styles_button.configure(state="disabled")
			self.sort_styles_button.configure(state="disabled")
			self.settings_button.configure(state="disabled")
			self.refresh_models_button.configure(state="disabled")
			self.refresh_prices_button.configure(state="disabled")
		elif self.prices_loading:
			self.start_button.configure(state="disabled")
			self.source_combo.configure(state="disabled")
			self.target_combo.configure(state="disabled")
			self.languages_button.configure(state="disabled")
			self.model_picker_button.configure(state="disabled")
			self._close_model_picker()
			self.settings_button.configure(state="disabled")
			self.refresh_models_button.configure(state="disabled")
			self.refresh_prices_button.configure(state="disabled")
		else:
			self.start_button.configure(state="normal")
			self.source_combo.configure(state="readonly")
			self.target_combo.configure(state="readonly")
			self.languages_button.configure(state="normal")
			self.model_picker_button.configure(state="normal")
			self.settings_button.configure(state="normal")
			self.refresh_models_button.configure(state="normal")
			self.refresh_prices_button.configure(state="normal")

	def _effective_api_key(self) -> str:
		saved_key = self.api_key_var.get().strip()
		if saved_key:
			return saved_key
		return os.getenv("OPENAI_API_KEY", "").strip()

	def _selected_model_id(self) -> str:
		return extract_model_name(self.model_var.get().strip())

	def _place_dialog_centered_on_monitor(
		self,
		win: tk.Toplevel,
		width: int,
		height: int,
		anchor_widget: tk.Misc | None = None,
	) -> None:
		# Keep dialogs consistently centered inside the main app window.
		self.root.update_idletasks()
		root_x = int(self.root.winfo_rootx())
		root_y = int(self.root.winfo_rooty())
		root_w = max(360, int(self.root.winfo_width()))
		root_h = max(280, int(self.root.winfo_height()))
		final_w = max(320, min(int(width), root_w - 16))
		final_h = max(220, min(int(height), root_h - 16))
		x = root_x + max(0, (root_w - final_w) // 2)
		y = root_y + max(0, (root_h - final_h) // 2)
		win.geometry(f"{final_w}x{final_h}+{x}+{y}")

	def _new_modal(self, title: str, width: int, height: int, min_width: int, min_height: int) -> tk.Toplevel:
		win = tk.Toplevel(self.root)
		win.title(title)
		self._place_dialog_centered_on_monitor(win, width, height)
		win.minsize(min_width, min_height)
		win.grab_set()
		return win

	def _resize_model_tree_columns(self, tree: ttk.Treeview, total_width: int | None = None) -> None:
		if tree is None:
			return
		if total_width is None or total_width <= 0:
			total_width = tree.winfo_width()
		if total_width <= 0:
			return
		# Keep model column dominant while three price columns share remaining width.
		available = max(total_width - 8, 520)
		model_w = max(200, int(available * 0.42))
		price_w = max(110, int((available - model_w) / 3))
		tree.column("model", width=model_w)
		tree.column("input", width=price_w)
		tree.column("cached", width=price_w)
		tree.column("output", width=price_w)

	def _schedule_model_selector_resize(self, width: int) -> None:
		if width <= 0 or not hasattr(self, "model_display_table"):
			return
		last_width = getattr(self, "_last_model_selector_resize_width", None)
		if last_width is not None and abs(last_width - width) < 12:
			return
		job = getattr(self, "_model_selector_resize_job", None)
		if job is not None:
			try:
				self.root.after_cancel(job)
			except tk.TclError:
				pass
		self._model_selector_resize_job = self.root.after(16, lambda: self._flush_model_selector_resize(width))

	def _flush_model_selector_resize(self, width: int) -> None:
		self._model_selector_resize_job = None
		self._last_model_selector_resize_width = width
		if hasattr(self, "model_display_table") and self.model_display_table.winfo_exists():
			self._resize_model_tree_columns(self.model_display_table, width)

	def _on_model_selector_resize(self, event) -> None:
		self._schedule_model_selector_resize(event.width - 36)

	def _model_output_usd(self, model_id: str) -> float | None:
		pricing = model_pricing_for(model_id)
		if pricing is None:
			return None
		return _coerce_float(pricing.get("output"))

	def _log_price_band_index(self, value: float, min_output: float, max_output: float, bands: int = 7) -> int:
		if bands <= 1 or max_output <= min_output:
			return 0
		# Use logarithmic scaling so very large price ranges still produce meaningful color separation.
		epsilon = 1e-9
		min_val = max(min_output, epsilon)
		max_val = max(max_output, min_val + epsilon)
		clamped = min(max(value, min_val), max_val)
		log_min = math.log(min_val)
		log_max = math.log(max_val)
		if log_max <= log_min:
			return 0
		ratio = (math.log(clamped) - log_min) / (log_max - log_min)
		idx = int(round(ratio * (bands - 1)))
		return max(0, min(bands - 1, idx))

	def _model_price_color(self, model_id: str, min_output: float, max_output: float) -> str:
		output = self._model_output_usd(model_id)
		if output is None:
			return "#ffe8cc"
		# 7 fixed pastel bands from green (cheap) to stronger red (expensive).
		palette = [
			"#dff3df",
			"#e8f5da",
			"#f2f5d8",
			"#f9f0d6",
			"#f9e5d0",
			"#f7d6cf",
			"#f2a6a6",
		]
		band_idx = self._log_price_band_index(float(output), min_output, max_output, bands=len(palette))
		return palette[band_idx]

	def _price_color_map_for_models(self) -> dict[str, str]:
		valid_values: list[float] = []
		for model_id in self.model_choices:
			value = self._model_output_usd(model_id)
			if value is not None:
				valid_values.append(float(value))
		if valid_values:
			min_output = min(valid_values)
			max_output = max(valid_values)
		else:
			min_output = 0.0
			max_output = 0.0
		return {
			model_id: self._model_price_color(model_id, min_output, max_output)
			for model_id in self.model_choices
		}

	def _apply_row_color_tag(
		self,
		tree: ttk.Treeview,
		item_id: str,
		color_hex: str,
		is_selected_model: bool = False,
	) -> None:
		tag = f"price_{color_hex.lstrip('#')}"
		tree.tag_configure(tag, background=color_hex, foreground="#1f2937")
		tags: list[str] = [tag]
		if is_selected_model:
			tree.tag_configure("selected_model_row", font=("Segoe UI", 10, "bold"), foreground="#111827")
			tags.append("selected_model_row")
		tree.item(item_id, tags=tuple(tags))

	def _refresh_model_display_row(self) -> None:
		if not hasattr(self, "model_display_table"):
			return
		table = self.model_display_table
		for item_id in table.get_children():
			table.delete(item_id)
		model_id = extract_model_name(self.model_var.get().strip())
		if not model_id and self.model_choices:
			model_id = self.model_choices[0]
		if not model_id:
			row_id = table.insert("", "end", values=("-", "-$ | -€", "-$ | -€", "-$ | -€"))
			self._apply_row_color_tag(table, row_id, "#ffe8cc")
			self.model_display_var.set("")
			self._resize_model_tree_columns(table)
			return
		input_text, cached_text, output_text = model_price_texts(model_id)
		price_colors = self._price_color_map_for_models()
		row_id = table.insert("", "end", values=(model_id, input_text, cached_text, output_text))
		self._apply_row_color_tag(table, row_id, price_colors.get(model_id, "#ffe8cc"))
		self.model_display_var.set(build_model_choice_label(model_id))
		self._resize_model_tree_columns(table)

	def _populate_model_picker_tree(self) -> None:
		if self.model_picker_tree is None:
			return
		selected_model = self._selected_model_id()
		for item_id in self.model_picker_tree.get_children():
			self.model_picker_tree.delete(item_id)
		price_colors = self._price_color_map_for_models()
		selected_item_id: str | None = None
		for idx, model_id in enumerate(self._sorted_model_choices_for_picker()):
			input_text, cached_text, output_text = model_price_texts(model_id)
			row_id = self.model_picker_tree.insert(
				"",
				"end",
				iid=f"m{idx}",
				values=(model_id, input_text, cached_text, output_text),
			)
			is_selected = bool(selected_model and model_id == selected_model)
			self._apply_row_color_tag(
				self.model_picker_tree,
				row_id,
				price_colors.get(model_id, "#ffe8cc"),
				is_selected_model=is_selected,
			)
			if is_selected:
				selected_item_id = row_id
		if selected_model:
			if selected_item_id:
				self.model_picker_tree.focus(selected_item_id)
				self.model_picker_tree.see(selected_item_id)
		self._resize_model_tree_columns(self.model_picker_tree)

	def _model_picker_usd_value(self, model_id: str, column: str) -> float | None:
		pricing = model_pricing_for(model_id)
		if pricing is None:
			return None
		field_map = {
			"input": "input",
			"cached": "cached_input",
			"output": "output",
		}
		field = field_map.get(column)
		if field is None:
			return None
		return _coerce_float(pricing.get(field))

	def _sorted_model_choices_for_picker(self) -> list[str]:
		models = list(self.model_choices)
		column = self.model_picker_sort_column
		desc = self.model_picker_sort_desc
		if column == "model":
			return sorted(models, key=lambda item: str(item).lower(), reverse=desc)

		with_price: list[str] = []
		without_price: list[str] = []
		for model_id in models:
			if self._model_picker_usd_value(model_id, column) is None:
				without_price.append(model_id)
			else:
				with_price.append(model_id)

		with_price.sort(
			key=lambda item: float(self._model_picker_usd_value(item, column) or 0.0),
			reverse=desc,
		)
		without_price.sort(key=lambda item: str(item).lower())
		return with_price + without_price

	def _set_model_picker_sort(self, column: str) -> None:
		if column == self.model_picker_sort_column:
			self.model_picker_sort_desc = not self.model_picker_sort_desc
		else:
			self.model_picker_sort_column = column
			self.model_picker_sort_desc = False
		self._refresh_model_picker_headings()
		self._populate_model_picker_tree()

	def _refresh_model_picker_headings(self) -> None:
		if self.model_picker_tree is None:
			return
		labels = {
			"model": "Model",
			"input": "Input",
			"cached": "Cached",
			"output": "Output",
		}
		for column, base_label in labels.items():
			label = base_label
			if column == self.model_picker_sort_column:
				label = f"{base_label} {'▼' if self.model_picker_sort_desc else '▲'}"
			self.model_picker_tree.heading(column, text=label)

	def _close_model_picker(self) -> None:
		if self.model_picker_window is not None:
			try:
				self.model_picker_window.destroy()
			except Exception:
				pass
		self.model_picker_window = None
		self.model_picker_tree = None

	def _open_model_picker_from_event(self, _event=None):
		if self.is_running or self.models_loading or self.prices_loading:
			return "break"
		self._open_model_picker()
		return "break"

	def _open_model_picker(self) -> None:
		if self.is_running or self.models_loading or self.prices_loading:
			return
		if not self.model_choices:
			self.status_var.set("No models loaded. Click 'Load Models' first.")
			return
		if self.model_picker_window is not None:
			try:
				self.model_picker_window.focus_set()
				return
			except Exception:
				self._close_model_picker()

		win = tk.Toplevel(self.root)
		self.model_picker_window = win
		win.title("Select OpenAI Model")
		win.transient(self.root)
		win.resizable(True, True)
		win.protocol("WM_DELETE_WINDOW", self._close_model_picker)

		self.root.update_idletasks()
		anchor = self.model_picker_button
		anchor_width = max(self.model_selector_frame.winfo_width(), 620)
		height = min(420, 80 + max(1, min(len(self.model_choices), 14)) * 24)

		# Keep the model picker strictly inside the main window bounds.
		root_x = int(self.root.winfo_rootx())
		root_y = int(self.root.winfo_rooty())
		root_w = max(360, int(self.root.winfo_width()))
		root_h = max(280, int(self.root.winfo_height()))
		margin = 8
		gap = 8
		final_w = max(320, min(int(anchor_width), max(320, root_w - (margin * 2))))
		final_h = max(220, min(int(height), max(220, root_h - (margin * 2))))

		anchor_rel_x = int(anchor.winfo_rootx()) - root_x
		anchor_rel_y = int(anchor.winfo_rooty()) - root_y
		anchor_h = int(anchor.winfo_height())
		anchor_w = int(anchor.winfo_width())

		x_rel = anchor_rel_x + anchor_w + gap
		if x_rel + final_w > root_w - margin:
			x_rel = anchor_rel_x - final_w - gap
		max_x_rel = max(margin, root_w - final_w - margin)
		x_rel = min(max(margin, x_rel), max_x_rel)

		y_rel = anchor_rel_y + (anchor_h // 2) - (final_h // 2)
		max_y_rel = max(margin, root_h - final_h - margin)
		y_rel = min(max(margin, y_rel), max_y_rel)

		win.geometry(f"{final_w}x{final_h}+{root_x + x_rel}+{root_y + y_rel}")

		wrap = ttk.Frame(win, padding=(8, 8))
		wrap.pack(fill="both", expand=True)
		wrap.columnconfigure(0, weight=1)
		wrap.rowconfigure(0, weight=1)

		tree = ttk.Treeview(
			wrap,
			columns=("model", "input", "cached", "output"),
			show="headings",
			selectmode="none",
		)
		tree.heading("model", text="Model", command=lambda: self._set_model_picker_sort("model"))
		tree.heading("input", text="Input", command=lambda: self._set_model_picker_sort("input"))
		tree.heading("cached", text="Cached", command=lambda: self._set_model_picker_sort("cached"))
		tree.heading("output", text="Output", command=lambda: self._set_model_picker_sort("output"))
		tree.column("model", anchor="w", width=300, stretch=True)
		tree.column("input", anchor="e", width=120, stretch=True)
		tree.column("cached", anchor="e", width=120, stretch=True)
		tree.column("output", anchor="e", width=120, stretch=True)
		tree.grid(row=0, column=0, sticky="nsew")
		scroll = ttk.Scrollbar(wrap, orient="vertical", command=tree.yview)
		scroll.grid(row=0, column=1, sticky="ns")
		tree.configure(yscrollcommand=scroll.set)
		tree.bind("<Configure>", lambda event: self._resize_model_tree_columns(tree, event.width))
		self.model_picker_tree = tree
		self._refresh_model_picker_headings()
		self._populate_model_picker_tree()

		def _choose_selection(item_id: str | None = None) -> None:
			if not item_id:
				item_id = tree.focus()
			if not item_id:
				return
			values = tree.item(item_id, "values")
			if not values:
				return
			model_id = extract_model_name(str(values[0]))
			if not model_id:
				return
			self.model_var.set(model_id)
			self._sync_model_selection_from_ui()
			self._persist_model_selection()
			self._close_model_picker()

		def _choose_click(event) -> None:
			region = tree.identify_region(event.x, event.y)
			if region not in {"tree", "cell"}:
				return
			clicked_item = tree.identify_row(event.y)
			if not clicked_item:
				return
			tree.focus(clicked_item)
			_choose_selection(clicked_item)

		def _choose_enter(_event=None) -> None:
			_choose_selection(tree.focus())

		tree.bind("<Double-1>", _choose_click)
		tree.bind("<Return>", _choose_enter)
		tree.bind("<ButtonRelease-1>", _choose_click)
		win.bind("<Escape>", lambda _event: self._close_model_picker())
		tree.focus_set()

	def _set_model_choices(self, model_ids: list[str]) -> None:
		self.model_choices = list(model_ids)
		self.model_choice_labels = [
			build_model_choice_label(model_id)
			for model_id in self.model_choices
		]
		self.model_label_to_id = dict(zip(self.model_choice_labels, self.model_choices))
		self.model_id_to_label = {model_id: label for label, model_id in self.model_label_to_id.items()}

		current_model = extract_model_name(self.model_var.get().strip())
		if self.model_choices and current_model not in self.model_choices:
			current_model = self.model_choices[0]
		if current_model:
			self.model_var.set(current_model)
			self.model_display_var.set(self.model_id_to_label.get(current_model, build_model_choice_label(current_model)))
		self._refresh_model_display_row()
		self._populate_model_picker_tree()

	def _sync_model_selection_from_ui(self) -> None:
		model_id = self._selected_model_id()
		if model_id:
			self.model_var.set(model_id)
			self.model_display_var.set(self.model_id_to_label.get(model_id, build_model_choice_label(model_id)))
		self._refresh_model_display_row()

	def _set_running(self, running: bool) -> None:
		self.is_running = running
		if not running:
			self.is_paused = False
			self.pause_event.clear()
		self._sync_controls()

	def _init_taskbar(self) -> None:
		try:
			if TaskbarProgress is None:
				self._taskbar = None
				return
			hwnd = self._taskbar_hwnd()
			self._taskbar = TaskbarProgress(hwnd) if hwnd > 0 else None
		except Exception:
			self._taskbar = None

	def _taskbar_update(self, *, percent: float | None = None, state: ProgressType | None = None) -> None:
		if self._taskbar is None:
			return
		try:
			if state is not None:
				self._taskbar.set_progress_type(state)
			if percent is not None:
				self._taskbar.set_progress(int(percent))
		except Exception:
			pass

	def _taskbar_hwnd(self) -> int:
		"""Return a stable top-level HWND for Tk windows (taskbar button owner)."""
		self.root.update_idletasks()
		try:
			frame = self.root.wm_frame()
			if isinstance(frame, str) and frame:
				# Tk returns hex window IDs like "0x123abc".
				if frame.lower().startswith("0x"):
					hwnd = int(frame, 16)
				else:
					hwnd = int(frame)
				if hwnd > 0:
					return hwnd
		except Exception:
			pass
		try:
			return int(self.root.winfo_id())
		except Exception:
			return 0

	def _taskbar_flash_done(self) -> None:
		if self._taskbar is None:
			return
		try:
			self._taskbar.flash_done()
		except Exception:
			pass

	def _refresh_counters(self) -> None:
		self.warning_counter_var.set(f"Warnings: {self.warning_count_live}")
		self.error_counter_var.set(f"Errors: {self.error_count_live}")

	def _refresh_usage_counters(self) -> None:
		self.token_counter_var.set(
			f"Tokens: In: {self.usage_prompt_tokens_live} | Out: {self.usage_completion_tokens_live}"
		)
		if self.usage_pricing_known_live:
			if self.usage_eur_rate_live is not None:
				self.cost_counter_var.set(
					f"Cost: {self.usage_cost_usd_live:.2f}$/{self.usage_cost_eur_live:.2f}€"
				)
			else:
				self.cost_counter_var.set(f"Cost: {self.usage_cost_usd_live:.2f}$/-€")
		else:
			self.cost_counter_var.set("Cost: n/a")

	def _set_models_loading(self, loading: bool) -> None:
		self.models_loading = loading
		self._sync_controls()

	def _set_prices_loading(self, loading: bool) -> None:
		self.prices_loading = loading
		self._sync_controls()

	def _request_model_list(self, on_startup: bool = False) -> None:
		api_key = self._effective_api_key()
		if not api_key:
			if on_startup:
				self.status_var.set("No API key set. Please configure it in Settings.")
			return

		if self.model_loader and self.model_loader.is_alive():
			return

		self.status_var.set("Checking available OpenAI models...")
		self._set_models_loading(True)
		if not self.is_running:
			self._taskbar_update(state=ProgressType.INDETERMINATE)

		self.model_loader = threading.Thread(
			target=self._run_model_discovery,
			args=(api_key, on_startup),
			daemon=True,
		)
		self.model_loader.start()

	def _run_model_discovery(self, api_key: str, on_startup: bool) -> None:
		try:
			model_ids = fetch_accessible_models(api_key)
			self.queue.put(("models_loaded", model_ids))
		except Exception as exc:
			details = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__))
			self.queue.put(("models_error", str(exc), details, on_startup))

	def _request_model_price_update(self) -> None:
		api_key = self._effective_api_key()
		if not api_key:
			self.status_var.set("No API key set. Please configure it in Settings.")
			return

		if self.model_price_loader and self.model_price_loader.is_alive():
			return

		model_ids = sorted({extract_model_name(model_id) for model_id in self.model_choices if extract_model_name(model_id)})
		if not model_ids:
			self.status_var.set("No local model list loaded. Resolving models from OpenAI...")

		self.progress_var.set(0)
		self.status_var.set("Updating model prices from OpenAI (best effort)...")
		self._set_prices_loading(True)
		if not self.is_running:
			# 0% is effectively invisible in the Windows taskbar overlay.
			self._taskbar_update(state=ProgressType.NORMAL, percent=1.0)
		self.model_price_loader = threading.Thread(
			target=self._run_model_price_update,
			args=(api_key, model_ids),
			daemon=True,
		)
		self.model_price_loader.start()

	def _run_model_price_update(self, api_key: str, model_ids: list[str]) -> None:
		try:
			def progress_update(done: int, total: int, message: str) -> None:
				self.queue.put(("prices_progress", done, total, message))

			summary = update_model_pricing_file(api_key, model_ids, progress_callback=progress_update)
			self.queue.put(("prices_updated", summary))
		except Exception as exc:
			details = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__))
			self.queue.put(("prices_error", str(exc), details))

	def _append_error_log(self, title: str, lines: list[str]) -> None:
		if self.error_log_text is None:
			return

		timestamp = datetime.now().strftime("%H:%M:%S")
		self.error_count += 1
		self.error_log_text.configure(state="normal")
		self.error_log_text.insert("end", f"[{self.error_count:03}] {title}\n", "entry_header")
		self.error_log_text.insert("end", f"Time: {timestamp}\n", "entry_meta")
		for line in lines:
			self.error_log_text.insert("end", f"{line}\n", "entry_body")
		self.error_log_text.insert("end", "-" * 76 + "\n", "entry_sep")
		self.error_log_text.see("end")
		self.error_log_text.configure(state="disabled")

	def _append_warning_log(self, title: str, lines: list[str]) -> None:
		if self.error_log_text is None:
			return

		timestamp = datetime.now().strftime("%H:%M:%S")
		self.error_count += 1
		self.error_log_text.configure(state="normal")
		self.error_log_text.insert("end", f"[{self.error_count:03}] {title}\n", "warn_header")
		self.error_log_text.insert("end", f"Time: {timestamp}\n", "warn_meta")
		for line in lines:
			self.error_log_text.insert("end", f"{line}\n", "entry_body")
		self.error_log_text.insert("end", "-" * 76 + "\n", "entry_sep")
		self.error_log_text.see("end")
		self.error_log_text.configure(state="disabled")

	def _append_info_log(self, title: str, lines: list[str]) -> None:
		if self.error_log_text is None:
			return

		timestamp = datetime.now().strftime("%H:%M:%S")
		self.error_count += 1
		self.error_log_text.configure(state="normal")
		self.error_log_text.insert("end", f"[{self.error_count:03}] {title}\n", "info_header")
		self.error_log_text.insert("end", f"Time: {timestamp}\n", "info_meta")
		for line in lines:
			self.error_log_text.insert("end", f"{line}\n", "entry_body")
		self.error_log_text.insert("end", "-" * 76 + "\n", "entry_sep")
		self.error_log_text.see("end")
		self.error_log_text.configure(state="disabled")

	def _clear_error_log(self) -> None:
		if self.error_log_text is None:
			return
		self.error_log_text.configure(state="normal")
		self.error_log_text.delete("1.0", "end")
		self.error_log_text.configure(state="disabled")

	def _update_styles_info(self) -> None:
		n_total = len(self.available_styles)
		if n_total == 0:
			text = "Select a file first" if not self.file_var.get().strip() else "Reading styles..."
			self.styles_info_var.set(text)
		else:
			n_selected = len(self.selected_styles & set(self.available_styles))
			self.styles_info_var.set(f"{n_selected} of {n_total} styles selected")

	def _update_char_styles_info(self) -> None:
		n_total = len(self.available_char_styles)
		if n_total == 0:
			text = "Select a file first" if not self.file_var.get().strip() else "No character styles found"
			self.char_styles_info_var.set(text)
		else:
			n_selected = len(self.selected_char_styles & set(self.available_char_styles))
			self.char_styles_info_var.set(f"{n_selected} of {n_total} character styles selected")

	def _current_glossary_pair_key(self) -> str:
		return language_pair_key(self.source_lang_var.get().strip(), self.target_lang_var.get().strip())

	def _apply_language_codes(self, codes: list[str]) -> None:
		self.language_codes = _clean_language_codes(codes)
		self.source_combo.configure(values=self.language_codes)
		self.target_combo.configure(values=self.language_codes)
		source = _canon_lang(self.source_lang_var.get())
		target = _canon_lang(self.target_lang_var.get())
		if source not in self.language_codes:
			source = self.language_codes[0]
		if target not in self.language_codes:
			target = self.language_codes[1] if len(self.language_codes) > 1 else self.language_codes[0]
		self.source_lang_var.set(source)
		self.target_lang_var.set(target)
		self._update_glossary_info()

	def _update_glossary_info(self) -> None:
		key = self._current_glossary_pair_key()
		count = len(self.glossary_pairs.get(key, []))
		src = self.source_lang_var.get().strip() or "?"
		tgt = self.target_lang_var.get().strip() or "?"
		self.glossary_info_var.set(f"Glossary {src} ↔ {tgt}: {count} entries")

	def _on_language_change(self, _event=None) -> None:
		self.source_lang_var.set(_canon_lang(self.source_lang_var.get()))
		self.target_lang_var.set(_canon_lang(self.target_lang_var.get()))
		self._update_glossary_info()
		self._persist_model_selection()

	def _open_language_manager(self) -> None:
		win = self._new_modal("Manage Language Codes", 520, 560, 460, 460)

		existing = list(self.language_codes)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)
		ttk.Label(outer, text="Language Codes", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Use IETF language codes (examples: de, en-US, es-419).",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		list_wrap = ttk.Frame(outer)
		list_wrap.pack(fill="both", expand=True)
		listbox = tk.Listbox(list_wrap, font=("Segoe UI", 10), activestyle="none")
		scroll = ttk.Scrollbar(list_wrap, orient="vertical", command=listbox.yview)
		listbox.configure(yscrollcommand=scroll.set)
		listbox.pack(side="left", fill="both", expand=True)
		scroll.pack(side="right", fill="y")

		entry_var = tk.StringVar()

		def refresh_list() -> None:
			listbox.delete(0, "end")
			for code in existing:
				listbox.insert("end", code)

		def on_select(_event=None) -> None:
			selection = listbox.curselection()
			if not selection:
				return
			idx = int(selection[0])
			if 0 <= idx < len(existing):
				entry_var.set(existing[idx])

		refresh_list()
		listbox.bind("<<ListboxSelect>>", on_select)

		entry_row = ttk.Frame(outer)
		entry_row.pack(fill="x", pady=(10, 0))
		ttk.Entry(entry_row, textvariable=entry_var, font=("Segoe UI", 10)).pack(side="left", fill="x", expand=True)

		btn_row = ttk.Frame(outer)
		btn_row.pack(fill="x", pady=(10, 0))

		def add_code() -> None:
			code = _canon_lang(entry_var.get())
			if not code or code in existing:
				return
			existing.append(code)
			entry_var.set("")
			refresh_list()

		def update_selected() -> None:
			selection = listbox.curselection()
			if not selection:
				return
			idx = int(selection[0])
			code = _canon_lang(entry_var.get())
			if not code:
				return
			if code in existing and existing[idx] != code:
				return
			existing[idx] = code
			refresh_list()
			listbox.selection_set(idx)

		def remove_selected() -> None:
			selection = listbox.curselection()
			if not selection or len(existing) <= 1:
				return
			idx = int(selection[0])
			del existing[idx]
			entry_var.set("")
			refresh_list()

		ttk.Button(btn_row, text="Add", style="Secondary.TButton", command=add_code).pack(side="left")
		ttk.Button(btn_row, text="Update selected", style="Secondary.TButton", command=update_selected).pack(side="left", padx=(8, 0))
		ttk.Button(btn_row, text="Delete selected", style="Secondary.TButton", command=remove_selected).pack(side="left", padx=(8, 0))

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))

		def apply_changes() -> None:
			self._apply_language_codes(existing)
			self._save_settings(language_codes=list(self.language_codes))
			self.status_var.set("Language codes updated.")
			win.destroy()

		ttk.Button(bottom, text="Apply", style="Primary.TButton", command=apply_changes).pack(side="right")
		ttk.Button(bottom, text="Cancel", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _open_glossary_editor(self) -> None:
		pair_key = self._current_glossary_pair_key()
		existing = [dict(item) for item in self.glossary_pairs.get(pair_key, [])]
		src = self.source_lang_var.get().strip() or "?"
		tgt = self.target_lang_var.get().strip() or "?"

		win = self._new_modal(f"Terminology Glossary: {src} ↔ {tgt}", 720, 620, 620, 520)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)
		ttk.Label(outer, text=f"Terminology table for {src} ↔ {tgt}", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Manage bidirectional word pairs. The AI should preferably use these pairs.",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		list_wrap = ttk.Frame(outer)
		list_wrap.pack(fill="both", expand=True)
		listbox = tk.Listbox(list_wrap, font=("Segoe UI", 10), activestyle="none")
		scroll = ttk.Scrollbar(list_wrap, orient="vertical", command=listbox.yview)
		listbox.configure(yscrollcommand=scroll.set)
		listbox.pack(side="left", fill="both", expand=True)
		scroll.pack(side="right", fill="y")

		def refresh_list() -> None:
			listbox.delete(0, "end")
			for item in existing:
				listbox.insert("end", f"{item['left']}  ↔  {item['right']}")

		refresh_list()

		entry_row = ttk.Frame(outer)
		entry_row.pack(fill="x", pady=(10, 0))
		left_var = tk.StringVar()
		right_var = tk.StringVar()
		ttk.Entry(entry_row, textvariable=left_var, font=("Segoe UI", 10)).pack(side="left", fill="x", expand=True)
		ttk.Label(entry_row, text="↔", style="Status.TLabel").pack(side="left", padx=8)
		ttk.Entry(entry_row, textvariable=right_var, font=("Segoe UI", 10)).pack(side="left", fill="x", expand=True)

		btn_row = ttk.Frame(outer)
		btn_row.pack(fill="x", pady=(10, 0))

		def add_pair() -> None:
			left = left_var.get().strip()
			right = right_var.get().strip()
			if not left or not right:
				return
			existing.append({"left": left, "right": right})
			left_var.set("")
			right_var.set("")
			refresh_list()

		def remove_selected() -> None:
			selection = listbox.curselection()
			if not selection:
				return
			idx = int(selection[0])
			if 0 <= idx < len(existing):
				del existing[idx]
			refresh_list()

		ttk.Button(btn_row, text="Add", style="Secondary.TButton", command=add_pair).pack(side="left")
		ttk.Button(btn_row, text="Delete selected", style="Secondary.TButton", command=remove_selected).pack(side="left", padx=(8, 0))

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))

		def apply_changes() -> None:
			cleaned: list[dict[str, str]] = []
			for item in existing:
				left = str(item.get("left", "")).strip()
				right = str(item.get("right", "")).strip()
				if left and right:
					cleaned.append({"left": left, "right": right})
			if cleaned:
				self.glossary_pairs[pair_key] = cleaned
			else:
				self.glossary_pairs.pop(pair_key, None)
			self._update_glossary_info()
			self._persist_model_selection()
			win.destroy()

		ttk.Button(bottom, text="Apply", style="Primary.TButton", command=apply_changes).pack(side="right")
		ttk.Button(bottom, text="Cancel", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _scan_styles_from_file(self, input_path: Path, on_startup: bool = False) -> None:
		self.styles_scanning = True
		self.styles_info_var.set("Reading styles...")
		self.char_styles_info_var.set("Reading character styles...")
		self._sync_controls()
		threading.Thread(target=self._run_style_scan, args=(input_path, on_startup), daemon=True).start()

	def _run_style_scan(self, input_path: Path, on_startup: bool) -> None:
		try:
			if is_docx_open_in_word(input_path):
				raise PermissionError("Datei ist in Word geoeffnet. Bitte schliessen und erneut versuchen.")
			if not zipfile.is_zipfile(input_path):
				raise ValueError("Datei ist kein gueltiges DOCX-Paket oder nicht lokal verfuegbar.")
			styles = collect_styles_from_docx(input_path)
			char_styles = collect_character_styles_from_docx(input_path)
			self.queue.put(("styles_scanned", styles, char_styles))
		except Exception as exc:
			self.queue.put(("styles_scan_error", str(exc), str(input_path), on_startup, is_docx_open_in_word(input_path)))

	def _open_style_selector(self) -> None:
		if not self.available_styles:
			return

		win = self._new_modal("Select Styles", 460, 520, 380, 400)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="Styles to translate", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Check the styles whose paragraphs should be translated.",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		btn_row = ttk.Frame(outer)
		btn_row.pack(fill="x", pady=(0, 8))
		check_vars: dict[str, tk.BooleanVar] = {}

		def select_all() -> None:
			for v in check_vars.values():
				v.set(True)

		def select_none() -> None:
			for v in check_vars.values():
				v.set(False)

		ttk.Button(btn_row, text="All", style="Secondary.TButton", command=select_all).pack(side="left")
		ttk.Button(btn_row, text="None", style="Secondary.TButton", command=select_none).pack(side="left", padx=(6, 0))

		list_outer = ttk.Frame(outer)
		list_outer.pack(fill="both", expand=True)

		canvas = tk.Canvas(list_outer, bg="#ffffff", highlightthickness=0)
		scrollbar = ttk.Scrollbar(list_outer, orient="vertical", command=canvas.yview)
		canvas.configure(yscrollcommand=scrollbar.set)
		scrollbar.pack(side="right", fill="y")
		canvas.pack(side="left", fill="both", expand=True)

		inner = ttk.Frame(canvas)
		canvas_window = canvas.create_window((0, 0), window=inner, anchor="nw")

		def _on_inner_configure(_event) -> None:
			canvas.configure(scrollregion=canvas.bbox("all"))

		def _on_canvas_configure(event) -> None:
			canvas.itemconfig(canvas_window, width=event.width)

		inner.bind("<Configure>", _on_inner_configure)
		canvas.bind("<Configure>", _on_canvas_configure)

		for style_name in self.available_styles:
			var = tk.BooleanVar(value=style_name in self.selected_styles)
			check_vars[style_name] = var
			ttk.Checkbutton(inner, text=style_name, variable=var).pack(anchor="w", padx=8, pady=3)

		def apply() -> None:
			self.selected_styles = {name for name, v in check_vars.items() if v.get()}
			self._update_styles_info()
			self._save_settings()
			win.destroy()

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))
		ttk.Button(bottom, text="Apply", style="Primary.TButton", command=apply).pack(side="right")
		ttk.Button(bottom, text="Cancel", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _open_char_style_selector(self) -> None:
		if not self.available_char_styles:
			return

		win = self._new_modal("Select Character Styles", 640, 620, 540, 480)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="Character styles to preserve", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Only selected character styles will be preserved via CHST anchors. All others will be ignored.",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		btn_row = ttk.Frame(outer)
		btn_row.pack(fill="x", pady=(0, 8))
		check_vars: dict[str, tk.BooleanVar] = {}

		def select_all() -> None:
			for v in check_vars.values():
				v.set(True)

		def select_none() -> None:
			for v in check_vars.values():
				v.set(False)

		ttk.Button(btn_row, text="All", style="Secondary.TButton", command=select_all).pack(side="left")
		ttk.Button(btn_row, text="None", style="Secondary.TButton", command=select_none).pack(side="left", padx=(6, 0))

		list_outer = ttk.Frame(outer)
		list_outer.pack(fill="both", expand=True)

		canvas = tk.Canvas(list_outer, bg="#ffffff", highlightthickness=0)
		scrollbar = ttk.Scrollbar(list_outer, orient="vertical", command=canvas.yview)
		canvas.configure(yscrollcommand=scrollbar.set)
		scrollbar.pack(side="right", fill="y")
		canvas.pack(side="left", fill="both", expand=True)

		inner = ttk.Frame(canvas)
		canvas_window = canvas.create_window((0, 0), window=inner, anchor="nw")

		def _on_inner_configure(_event) -> None:
			canvas.configure(scrollregion=canvas.bbox("all"))

		def _on_canvas_configure(event) -> None:
			canvas.itemconfig(canvas_window, width=event.width)

		inner.bind("<Configure>", _on_inner_configure)
		canvas.bind("<Configure>", _on_canvas_configure)

		for style_name in self.available_char_styles:
			var = tk.BooleanVar(value=style_name in self.selected_char_styles)
			check_vars[style_name] = var
			ttk.Checkbutton(inner, text=style_name, variable=var).pack(anchor="w", padx=8, pady=3)

		def apply() -> None:
			self.selected_char_styles = {name for name, v in check_vars.items() if v.get()}
			self._update_char_styles_info()
			self._save_settings()
			win.destroy()

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))
		ttk.Button(bottom, text="Apply", style="Primary.TButton", command=apply).pack(side="right")
		ttk.Button(bottom, text="Cancel", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _open_style_cleaner(self) -> None:
		fpath = Path(self.file_var.get().strip())
		if not fpath.exists():
			return

		all_para, all_char = collect_all_defined_styles_from_docx(fpath)
		used_para = set(self.available_styles)
		used_char = set(collect_character_styles_from_docx(fpath, include_named_bold_styles=True))
		unused_para = [s for s in all_para if s not in used_para]
		unused_char = [s for s in all_char if s not in used_char]

		win = self._new_modal("Clean Up Styles", 640, 560, 540, 440)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="Unused Styles", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Defined but not used in the document. Multi-select is supported.",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		cols = ttk.Frame(outer)
		cols.pack(fill="both", expand=True)
		cols.columnconfigure(0, weight=1)
		cols.columnconfigure(1, weight=1)
		cols.rowconfigure(1, weight=1)

		para_lbl = ttk.Label(cols, text=f"Paragraph Styles ({len(unused_para)})", font=("Segoe UI Semibold", 10))
		para_lbl.grid(row=0, column=0, sticky="w", padx=(0, 8), pady=(0, 4))

		char_lbl = ttk.Label(cols, text=f"Character Styles ({len(unused_char)})", font=("Segoe UI Semibold", 10))
		char_lbl.grid(row=0, column=1, sticky="w", pady=(0, 4))

		pf = ttk.Frame(cols)
		pf.grid(row=1, column=0, sticky="nsew", padx=(0, 8))
		pf.rowconfigure(0, weight=1)
		pf.columnconfigure(0, weight=1)
		para_lb = tk.Listbox(pf, selectmode="extended", font=("Segoe UI", 9), activestyle="none")
		para_lb.grid(row=0, column=0, sticky="nsew")
		psb = ttk.Scrollbar(pf, orient="vertical", command=para_lb.yview)
		psb.grid(row=0, column=1, sticky="ns")
		para_lb.configure(yscrollcommand=psb.set)
		for s in unused_para:
			para_lb.insert("end", s)

		cf = ttk.Frame(cols)
		cf.grid(row=1, column=1, sticky="nsew")
		cf.rowconfigure(0, weight=1)
		cf.columnconfigure(0, weight=1)
		char_lb = tk.Listbox(cf, selectmode="extended", font=("Segoe UI", 9), activestyle="none")
		char_lb.grid(row=0, column=0, sticky="nsew")
		csb = ttk.Scrollbar(cf, orient="vertical", command=char_lb.yview)
		csb.grid(row=0, column=1, sticky="ns")
		char_lb.configure(yscrollcommand=csb.set)
		for s in unused_char:
			char_lb.insert("end", s)

		hint = ttk.Label(outer, text="", style="Sub.TLabel")
		hint.pack(anchor="w", pady=(8, 0))

		def _remove_selected() -> None:
			pidxs = list(para_lb.curselection())
			cidxs = list(char_lb.curselection())
			if not pidxs and not cidxs:
				hint.configure(text="Please select at least one style.")
				return
			sel_para = [unused_para[i] for i in pidxs]
			sel_char = [unused_char[i] for i in cidxs]
			n = remove_styles_from_docx(fpath, sel_para, sel_char)
			for i in sorted(pidxs, reverse=True):
				para_lb.delete(i)
				unused_para.pop(i)
			for i in sorted(cidxs, reverse=True):
				char_lb.delete(i)
				unused_char.pop(i)
			para_lbl.configure(text=f"Paragraph Styles ({len(unused_para)})")
			char_lbl.configure(text=f"Character Styles ({len(unused_char)})")
			hint.configure(text=f"{n} style(s) removed.")

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))
		ttk.Button(bottom, text="Remove from File", style="Primary.TButton", command=_remove_selected).pack(side="right")
		ttk.Button(bottom, text="Close", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _open_style_priority_editor(self) -> None:
		fpath = Path(self.file_var.get().strip())
		if not fpath.exists():
			return

		items = collect_all_styles_with_priority_from_docx(fpath)
		if not items:
			return

		win = self._new_modal("Sort Styles", 940, 640, 820, 520)

		outer = ttk.Frame(win, padding=16)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="Styles by Priority", font=("Segoe UI Semibold", 12)).pack(anchor="w")
		ttk.Label(
			outer,
			text="Double-click on priority to edit. Ctrl/Shift+Click for multi-select. Drag & Drop: drag selected entries to a new position.",
			font=("Segoe UI", 9),
			foreground="#445069",
		).pack(anchor="w", pady=(2, 10))

		table_wrap = ttk.Frame(outer)
		table_wrap.pack(fill="both", expand=True)
		table_wrap.columnconfigure(0, weight=1)
		table_wrap.rowconfigure(0, weight=1)

		cols = ("prio", "typ", "name", "sid")
		tree = ttk.Treeview(
			table_wrap,
			columns=cols,
			show="headings",
			selectmode="extended",
			height=22,
		)
		tree.heading("prio", text="Priority")
		tree.heading("typ", text="Type")
		tree.heading("name", text="Name")
		tree.heading("sid", text="Style-ID")
		tree.column("prio", width=90, anchor="center", stretch=False)
		tree.column("typ", width=100, anchor="w", stretch=False)
		tree.column("name", width=320, anchor="w", stretch=True)
		tree.column("sid", width=180, anchor="w", stretch=True)
		tree.grid(row=0, column=0, sticky="nsew")

		scroll = ttk.Scrollbar(table_wrap, orient="vertical", command=tree.yview)
		scroll.grid(row=0, column=1, sticky="ns")
		tree.configure(yscrollcommand=scroll.set)

		hint = ttk.Label(outer, text="", style="Sub.TLabel")
		hint.pack(anchor="w", pady=(8, 0))

		# Drop-indicator line drawn directly over the Treeview widget
		indicator_line = tk.Frame(tree, bg="#1a6ecf", height=2)

		drag_state: dict = {
			"active": False,   # True once mouse moved past threshold
			"press_x": 0,
			"press_y": 0,
			"start_row": None,
			"selection": [],   # rows captured when drag starts
			"ghost": None,     # Toplevel ghost window
		}
		edit_state: dict = {
			"entry": None,
			"row": None,
		}
		_DRAG_THRESHOLD = 5

		def _sort_items() -> None:
			items.sort(key=lambda item: (int(item["priority"]), str(item["name"]).lower(), str(item["style_id"]).lower()))

		def _refresh_tree(select_keys: set[str] | None = None) -> None:
			tree.delete(*tree.get_children())
			for idx, item in enumerate(items):
				tree.insert(
					"",
					"end",
					iid=str(idx),
					values=(
						str(item["priority"]),
						str(item["type"]),
						str(item["name"]),
						str(item["style_id"]),
					),
				)
			if select_keys:
				to_select: list[str] = []
				for idx, item in enumerate(items):
					if str(item["key"]) in select_keys:
						to_select.append(str(idx))
				if to_select:
					tree.selection_set(tuple(to_select))
					tree.see(to_select[0])

		_sort_items()
		_refresh_tree()

		def _hide_indicator() -> None:
			indicator_line.place_forget()

		def _show_indicator(y: int, target_iid: str) -> None:
			if not target_iid or not target_iid.isdigit():
				_hide_indicator()
				return
			bbox = tree.bbox(target_iid)
			if not bbox:
				_hide_indicator()
				return
			_, by, _, bh = bbox
			line_y = by if y <= by + bh // 2 else by + bh
			indicator_line.place(x=0, y=line_y, relwidth=1.0, height=2)
			indicator_line.lift()

		def _destroy_ghost() -> None:
			if drag_state["ghost"] is not None:
				try:
					drag_state["ghost"].destroy()
				except Exception:
					pass
				drag_state["ghost"] = None

		def _close_editor() -> None:
			if edit_state["entry"] is not None:
				edit_state["entry"].destroy()
				edit_state["entry"] = None
				edit_state["row"] = None

		def _begin_priority_edit(row_iid: str) -> None:
			_close_editor()
			bbox = tree.bbox(row_iid, "prio")
			if not bbox:
				return
			x, y, w, h = bbox
			entry = ttk.Entry(tree)
			entry.place(x=x, y=y, width=w, height=h)
			entry.insert(0, tree.set(row_iid, "prio"))
			entry.focus_set()
			entry.selection_range(0, "end")

			def _commit(_event=None) -> None:
				val_raw = entry.get().strip()
				if not val_raw.isdigit() or int(val_raw) < 1:
					hint.configure(text="Prioritaet muss eine ganze Zahl >= 1 sein.")
					entry.focus_set()
					return
				old_selection = {str(items[int(iid)]["key"]) for iid in tree.selection() if iid.isdigit()}
				items[int(row_iid)]["priority"] = int(val_raw)
				_sort_items()
				_refresh_tree(old_selection)
				hint.configure(text="Prioritaet geaendert.")
				_close_editor()

			def _cancel(_event=None) -> None:
				_close_editor()

			entry.bind("<Return>", _commit)
			entry.bind("<Escape>", _cancel)
			entry.bind("<FocusOut>", _commit)
			edit_state["entry"] = entry
			edit_state["row"] = row_iid

		def _on_double_click(event) -> None:
			row_iid = tree.identify_row(event.y)
			col = tree.identify_column(event.x)
			if not row_iid or col != "#1":
				return
			_begin_priority_edit(row_iid)

		def _on_press(event) -> None:
			_close_editor()
			_hide_indicator()
			row_iid = tree.identify_row(event.y)
			drag_state["start_row"] = row_iid
			drag_state["press_x"] = event.x
			drag_state["press_y"] = event.y
			drag_state["active"] = False
			drag_state["selection"] = []
			if not row_iid:
				return
			# Plain click on a row that is already inside a multi-selection:
			# return "break" so Tk does not collapse the selection down to one row.
			_MOD = 0x0004 | 0x0001  # Shift | Ctrl
			current_sel = list(tree.selection())
			if row_iid in current_sel and len(current_sel) > 1 and not (event.state & _MOD):
				drag_state["selection"] = current_sel
				return "break"

		def _on_motion(event) -> None:
			if drag_state["start_row"] is None:
				return
			dx = event.x - drag_state["press_x"]
			dy = event.y - drag_state["press_y"]
			if (dx * dx + dy * dy) ** 0.5 < _DRAG_THRESHOLD:
				return
			if not drag_state["active"]:
				drag_state["active"] = True
				sel = drag_state["selection"] or list(tree.selection())
				if not sel and drag_state["start_row"]:
					sel = [drag_state["start_row"]]
				drag_state["selection"] = sel
				# Build ghost window
				count = len(sel)
				first_iid = sel[0]
				first_name = str(items[int(first_iid)]["name"]) if first_iid.isdigit() and int(first_iid) < len(items) else "?"
				label = first_name if count == 1 else f"{first_name}  (+{count - 1} more)"
				g = tk.Toplevel(win)
				g.overrideredirect(True)
				g.attributes("-alpha", 0.78)
				g.attributes("-topmost", True)
				g.configure(bg="#dbe8ff")
				tk.Label(g, text=f"  {label}  ", bg="#dbe8ff", fg="#16315f",
						 font=("Segoe UI", 9), pady=4).pack()
				drag_state["ghost"] = g
			if drag_state["ghost"]:
				drag_state["ghost"].geometry(f"+{event.x_root + 14}+{event.y_root + 8}")
			_show_indicator(event.y, tree.identify_row(event.y))

		def _on_release(event) -> None:
			_hide_indicator()
			_destroy_ghost()
			if not drag_state["active"]:
				drag_state["start_row"] = None
				return
			drag_state["active"] = False
			drag_state["start_row"] = None
			selection_iids = [iid for iid in drag_state["selection"] if iid.isdigit()]
			drag_state["selection"] = []
			if not selection_iids:
				return
			target_iid = tree.identify_row(event.y)
			if not target_iid or not target_iid.isdigit():
				return
			sel_indexes = sorted(int(iid) for iid in selection_iids)
			# Determine insert position: above or below row centre
			insert_index = int(target_iid)
			bbox = tree.bbox(target_iid)
			if bbox and event.y > (bbox[1] + bbox[3] // 2):
				insert_index += 1
			moved = [items[idx] for idx in sel_indexes]
			remaining = [item for idx, item in enumerate(items) if idx not in set(sel_indexes)]
			shift = sum(1 for idx in sel_indexes if idx < insert_index)
			insert_index = max(0, min(len(remaining), insert_index - shift))
			for offset, item in enumerate(moved):
				remaining.insert(insert_index + offset, item)
			if remaining == items:
				return
			items[:] = remaining
			# All moved items get priority = (item above insertion point) + 1
			prio_above = int(items[insert_index - 1]["priority"]) if insert_index > 0 else 0
			new_prio = prio_above + 1
			for offset in range(len(moved)):
				items[insert_index + offset]["priority"] = new_prio
			# Cascade by contiguous priority groups so equal runs remain equal.
			cascade_prio = new_prio
			i = insert_index + len(moved)
			while i < len(items):
				group_prio = int(items[i]["priority"])
				group_end = i + 1
				while group_end < len(items) and int(items[group_end]["priority"]) == group_prio:
					group_end += 1
				if group_prio <= cascade_prio:
					cascade_prio += 1
					for j in range(i, group_end):
						items[j]["priority"] = cascade_prio
				else:
					cascade_prio = group_prio
				i = group_end
			selected_keys = {str(item["key"]) for item in moved}
			_sort_items()
			_refresh_tree(selected_keys)
			hint.configure(text=f"Priority {new_prio} set on {len(moved)} style(s).")

		tree.bind("<Double-1>", _on_double_click)
		tree.bind("<ButtonPress-1>", _on_press)
		tree.bind("<B1-Motion>", _on_motion)
		tree.bind("<ButtonRelease-1>", _on_release)

		def _save_changes() -> None:
			_close_editor()
			style_priority_by_key: dict[str, int] = {}
			for item in items:
				style_priority_by_key[str(item["key"])] = int(item["priority"])
			changed = set_style_priorities_in_docx(fpath, style_priority_by_key)
			self.status_var.set(f"Priorities saved ({changed} changes).")
			hint.configure(text=f"Saved: {changed} style(s) updated.")

		def _resort_by_priority() -> None:
			_close_editor()
			selected_keys = {str(items[int(iid)]["key"]) for iid in tree.selection() if iid.isdigit()}
			_sort_items()
			_refresh_tree(selected_keys)
			hint.configure(text="Sorted by priority.")

		bottom = ttk.Frame(outer)
		bottom.pack(fill="x", pady=(12, 0))
		ttk.Button(bottom, text="Sort by Priority", style="Secondary.TButton", command=_resort_by_priority).pack(side="left")
		ttk.Button(bottom, text="Save", style="Primary.TButton", command=_save_changes).pack(side="right")
		ttk.Button(bottom, text="Close", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=(0, 8))

	def _open_settings_window(self) -> None:
		win = self._new_modal("Settings", 760, 560, 700, 500)
		win.transient(self.root)

		outer = ttk.Frame(win, padding=GS)
		outer.pack(fill="both", expand=True)

		ttk.Label(outer, text="App Settings", style="Header.TLabel").pack(anchor="w")
		ttk.Label(
			outer,
			text="This window can be extended with additional options.",
			style="Sub.TLabel",
		).pack(anchor="w", pady=GS)

		openai_card = ttk.Frame(outer, style="Card.TFrame", padding=GS)
		openai_card.pack(fill="x")
		ttk.Label(openai_card, text="OpenAI", style="Label.TLabel").grid(row=0, column=0, columnspan=2, sticky="w")
		ttk.Label(openai_card, text="API Key", style="Status.TLabel").grid(row=1, column=0, sticky="w", pady=GS)

		api_key_edit_var = tk.StringVar(value=self.api_key_var.get())
		api_entry = ttk.Entry(openai_card, textvariable=api_key_edit_var, show="*", font=("Segoe UI", 10))
		api_entry.grid(row=2, column=0, columnspan=2, sticky="ew", pady=GS)

		key_hint = ttk.Label(
			openai_card,
			text="Leave empty to use OPENAI_API_KEY from the environment.",
			style="Sub.TLabel",
		)
		key_hint.grid(row=3, column=0, columnspan=2, sticky="w")

		show_var = tk.BooleanVar(value=False)

		def toggle_show() -> None:
			api_entry.configure(show="" if show_var.get() else "*")

		ttk.Checkbutton(
			openai_card,
			text="Show API Key",
			variable=show_var,
			command=toggle_show,
		).grid(row=4, column=0, sticky="w", pady=GS)

		openai_card.columnconfigure(0, weight=1)
		openai_card.columnconfigure(1, weight=0)

		future_card = ttk.Frame(outer, style="Card.TFrame", padding=GS)
		future_card.pack(fill="x", pady=GS)
		ttk.Label(future_card, text="Appearance", style="Label.TLabel").grid(row=0, column=0, sticky="w")
		ttk.Label(future_card, text="Theme", style="Status.TLabel").grid(row=1, column=0, sticky="w", pady=GS)

		theme_edit_var = tk.StringVar(value=_normalize_theme_mode(self.theme_var.get()))
		theme_combo = ttk.Combobox(
			future_card,
			textvariable=theme_edit_var,
			values=list(THEME_MODE_OPTIONS),
			state="readonly",
			font=("Segoe UI", 10),
		)
		theme_combo.grid(row=2, column=0, sticky="w")
		ttk.Label(
			future_card,
			text="System follows your OS theme.",
			style="Sub.TLabel",
		).grid(row=3, column=0, sticky="w", pady=GS)
		future_card.columnconfigure(0, weight=1)

		def apply_settings() -> None:
			new_api_key = api_key_edit_var.get().strip()
			new_theme_mode = _normalize_theme_mode(theme_edit_var.get())
			self.api_key_var.set(new_api_key)
			self.theme_var.set(new_theme_mode)
			self._apply_theme()
			self._save_settings(api_key=new_api_key, theme_mode=new_theme_mode)
			self.status_var.set("Settings saved.")
			self._request_model_list()
			win.destroy()

		buttons = ttk.Frame(outer)
		buttons.pack(fill="x", pady=GS)
		ttk.Button(buttons, text="Save", style="Primary.TButton", command=apply_settings).pack(side="right")
		ttk.Button(buttons, text="Cancel", style="Secondary.TButton", command=win.destroy).pack(side="right", padx=GS)

	def _open_output_folder(self) -> None:
		if self.last_output and self.last_output.exists():
			import subprocess
			subprocess.run(["explorer", "/select,", str(self.last_output)], check=False)

	def _cancel(self) -> None:
		if self.worker and self.worker.is_alive():
			self.cancel_event.set()
			self.pause_event.clear()
			self.status_var.set("Cancellation requested... waiting for current API request to finish.")
			self.cancel_button.configure(state="disabled")
			self.pause_button.configure(state="disabled")

	def _toggle_pause(self) -> None:
		if not (self.worker and self.worker.is_alive()):
			return
		if self.is_paused:
			self.pause_event.clear()
			self.status_var.set("Resuming translation...")
			self.pause_button.configure(state="disabled")
		else:
			self.pause_event.set()
			self.status_var.set("Pause requested... waiting for current block to finish.")
			self.pause_button.configure(state="disabled")

	def _format_status_lines(self, col1: str, col2: str = "", col3: str = "", col4: str = "") -> str:
		line1 = "  ".join(p for p in [col1, col2, col3] if p)
		return (line1 + "\n" + col4) if col4 else line1

	def _start(self) -> None:
		self._sync_model_selection_from_ui()

		if self.models_loading:
			self.status_var.set("Please wait: models are still being loaded from OpenAI.")
			self._append_error_log(
				"Start blocked",
				["Models are still loading. Please wait a moment and try again."],
			)
			return

		input_text = self.file_var.get().strip()
		model = self.model_var.get().strip()
		source_language = _canon_lang(self.source_lang_var.get().strip()) or "de"
		target_language = _canon_lang(self.target_lang_var.get().strip()) or "en-US"
		api_key = self._effective_api_key()
		selected_styles = frozenset(self.selected_styles)
		selected_char_styles = frozenset(self.selected_char_styles)

		if not input_text:
			self.status_var.set("Error: No input DOCX selected.")
			self._append_error_log("Validation Error", ["Please select an input DOCX file first."])
			return

		input_path = Path(input_text)
		if not input_path.exists() or input_path.suffix.lower() != ".docx":
			self.status_var.set("Error: Invalid input file.")
			self._append_error_log("Validation Error", ["Please select a valid .docx file."])
			return

		if is_docx_open_in_word(input_path):
			self.status_var.set("Error: File is open in Word.")
			self._append_error_log(
				"Validation Error",
				[
					f"File: {input_path}",
					"The file is currently open in Word (lock file detected).",
					"Please close the file in Word and try again.",
				],
			)
			return

		if not api_key:
			self.status_var.set("Error: API key missing.")
			self._append_error_log(
				"Validation Error",
				["OpenAI API key is missing. Please set it in Settings or via the OPENAI_API_KEY environment variable."],
			)
			return

		if not model:
			self.status_var.set("Error: No model selected.")
			self._append_error_log(
				"Validation Error",
				[
					"No model is selected.",
					"Please click 'Load Models' first and choose one model.",
				],
			)
			return

		if not selected_styles:
			self.status_var.set("Error: No style selected.")
			self._append_error_log(
				"Validation Error",
				["Please select at least one paragraph style (click 'Styles')."],
			)
			return

		if source_language == target_language:
			self.status_var.set("Error: Source and target language are the same.")
			self._append_error_log(
				"Validation Error",
				["Please select different source and target languages."],
			)
			return

		self._save_settings(input_file=input_text)

		self.last_output = build_output_path(input_path, target_language)
		self.last_output_exists = False
		self.reveal_button.grid_remove()
		self.warning_count_live = 0
		self.error_count_live = 0
		self._refresh_counters()
		self.usage_prompt_tokens_live = 0
		self.usage_completion_tokens_live = 0
		self.usage_total_tokens_live = 0
		self.usage_cost_usd_live = 0.0
		self.usage_cost_eur_live = 0.0
		self.usage_eur_rate_live = None
		self.usage_pricing_known_live = False
		self.api_request_count_live = 0
		self.last_request_id_live = "-"
		self.last_response_id_live = "-"
		self.current_block_label = ""
		self.current_block_start_tokens = 0
		self.current_block_start_cost_usd = 0.0
		self.current_block_start_cost_eur = 0.0
		self.live_info_var.set("Live-Info: -")
		self.api_request_info_var.set("Req-ID: - | Resp-ID: -")
		self._refresh_usage_counters()
		self.progress_var.set(0)
		self.status_var.set("Starting translation...")
		self.cancel_event.clear()
		self.pause_event.clear()
		self.is_paused = False
		self._taskbar_update(state=ProgressType.NORMAL)
		self._set_running(True)

		self.worker = threading.Thread(
			target=self._run_translation,
			args=(
				input_path,
				self.last_output,
				model,
				api_key,
				selected_styles,
				selected_char_styles,
				source_language,
				target_language,
				self.debug_var.get(),
			),
			daemon=True,
		)
		self.worker.start()

	def _persist_model_selection(self, _event=None) -> None:
		self._sync_model_selection_from_ui()
		self._save_settings()

	def _run_translation(
		self,
		input_path: Path,
		output_path: Path,
		model: str,
		api_key: str,
		selected_styles: frozenset[str],
		selected_char_styles: frozenset[str],
		source_language: str,
		target_language: str,
		debug_mode: bool,
	) -> None:
		try:
			active_glossary = get_glossary_for_pair(self.glossary_pairs, source_language, target_language)

			def usage_update(usage_summary: dict[str, object]) -> None:
				self.queue.put(("usage", usage_summary))

			def request_update(request_info: dict[str, object]) -> None:
				self.queue.put(("api_request", request_info))

			def pause_update(paused: bool, processed: int, total: int) -> None:
				self.queue.put(("paused_state", paused, processed, total, str(output_path)))

			translator = OpenAIParagraphTranslator(
				api_key=api_key,
				model=model,
				source_lang=source_language,
				target_lang=target_language,
				glossary_pairs=active_glossary,
				usage_callback=usage_update,
				request_callback=request_update,
			)

			def progress_update(
				processed: int,
				total: int,
				translated: int,
				failed: int,
				error_obj: Exception | None,
				paragraph_info: str,
				attempt: int,
				is_attempt: bool,
				debug_payload: dict | None,
			) -> None:
				self.queue.put(
					(
						"progress",
						processed,
						total,
						translated,
						failed,
						error_obj,
						paragraph_info,
						attempt,
						is_attempt,
						debug_payload,
					)
				)

			total, translated, failed, warning_count, cancelled, failed_items = translate_docx(
				input_path=input_path,
				output_path=output_path,
				translator=translator,
				selected_styles=selected_styles,
				selected_char_styles=set(selected_char_styles),
				target_language=target_language,
				debug_mode=debug_mode,
				progress_callback=progress_update,
				cancel_event=self.cancel_event,
				pause_event=self.pause_event,
				pause_callback=pause_update,
			)
			usage_summary = translator.get_usage_summary()
			self.queue.put(("done", total, translated, failed, warning_count, output_path, cancelled, failed_items, usage_summary))
		except Exception as exc:
			details = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__))
			self.queue.put(("error", exc, details))

	def _poll_queue(self) -> None:
		try:
			while True:
				item = self.queue.get_nowait()
				event = item[0]

				if event == "usage":
					_, usage_summary = item
					if isinstance(usage_summary, dict):
						self.usage_prompt_tokens_live = int(usage_summary.get("prompt_tokens", 0) or 0)
						self.usage_completion_tokens_live = int(usage_summary.get("completion_tokens", 0) or 0)
						self.usage_total_tokens_live = int(usage_summary.get("total_tokens", 0) or 0)
						self.usage_pricing_known_live = bool(usage_summary.get("pricing_known", False))
						estimated_usd = usage_summary.get("estimated_cost_usd")
						if self.usage_pricing_known_live and isinstance(estimated_usd, (int, float)):
							self.usage_cost_usd_live = float(estimated_usd)
						self.usage_eur_rate_live = _coerce_float(usage_summary.get("eur_rate"))
						estimated_eur = usage_summary.get("estimated_cost_eur")
						if self.usage_pricing_known_live and isinstance(estimated_eur, (int, float)):
							self.usage_cost_eur_live = float(estimated_eur)
						self._refresh_usage_counters()
					continue

				if event == "api_request":
					_, request_info = item
					if isinstance(request_info, dict):
						self.api_request_count_live = int(request_info.get("api_calls", self.api_request_count_live) or self.api_request_count_live)
						self.last_request_id_live = str(request_info.get("request_id", "-") or "-")
						self.last_response_id_live = str(request_info.get("response_id", "-") or "-")
						self.api_request_info_var.set(
							f"Req-ID: {self.last_request_id_live} | Resp-ID: {self.last_response_id_live}"
						)
					continue

				if event == "progress":
					_, processed, total, translated, failed, error_obj, paragraph_info, attempt, is_attempt, debug_payload = item
					percent = 100.0 if total == 0 else (processed / total) * 100.0
					self.progress_var.set(percent)
					self._taskbar_update(percent=percent)
					if is_attempt:
						is_precheck = paragraph_info == "Precheck"
						if attempt == 1 and paragraph_info and not is_precheck:
							self.current_block_label = str(paragraph_info)
							self.current_block_start_tokens = int(self.usage_total_tokens_live)
							self.current_block_start_cost_usd = float(self.usage_cost_usd_live)
							self.current_block_start_cost_eur = float(self.usage_cost_eur_live)
						current_idx = min(processed + 1, total) if total else 0
						if is_precheck:
							self.status_var.set(
								self._format_status_lines(f"Precheck... {processed}/{total}")
							)
						else:
							self.status_var.set(
								self._format_status_lines(
									"Translating...",
									f"{current_idx}/{total}",
									f"Attempt {attempt}",
									paragraph_info,
								)
							)
					elif isinstance(error_obj, HardAnchorFallbackWarning):
						self.warning_count_live += 1
						self._refresh_counters()
						warn_lines = [
							f"Paragraph: {paragraph_info}",
							f"Attempts: {attempt if attempt > 0 else '-'}",
							"Inline formulas could not be exactly restored at their anchor positions.",
							"The paragraph was translated anyway; formulas were placed approximately at the end.",
						]
						if getattr(error_obj, "diagnostics", None):
							warn_lines.append("--- Hard-Anchor Diagnostics ---")
							warn_lines.extend(list(error_obj.diagnostics))
						self._append_warning_log(
							"Warning: Formula position approximated",
							warn_lines,
						)
						self.status_var.set(
							self._format_status_lines(
								"Block with Warning",
								f"{processed}/{total}",
								f"Final errors: {failed}",
								paragraph_info,
							)
						)
					elif error_obj:
						self.error_count_live += 1
						self._refresh_counters()
						summary, detail_lines = format_openai_exception(error_obj)
						self._append_error_log(
							"Final Paragraph Error",
							[
								f"Paragraph: {paragraph_info}",
								f"Attempts: {attempt if attempt > 0 else '-'}",
								*detail_lines,
							],
						)
						self.status_var.set(
							self._format_status_lines(
								"Block failed",
								f"{processed}/{total}",
								f"Final errors: {failed}",
								paragraph_info,
							)
						)
					else:
						if self.current_block_label:
							delta_tokens = int(self.usage_total_tokens_live - self.current_block_start_tokens)
							if self.usage_pricing_known_live and self.usage_eur_rate_live is not None:
								delta_cost_usd = float(self.usage_cost_usd_live - self.current_block_start_cost_usd)
								delta_cost_eur = float(self.usage_cost_eur_live - self.current_block_start_cost_eur)
								self.live_info_var.set(
									f"Previous block: {self.current_block_label} | +{delta_tokens} tokens | +{delta_cost_usd:.2f}$ | +{delta_cost_eur:.2f}€"
								)
							elif self.usage_pricing_known_live:
								delta_cost_usd = float(self.usage_cost_usd_live - self.current_block_start_cost_usd)
								self.live_info_var.set(
									f"Previous block: {self.current_block_label} | +{delta_tokens} tokens | +{delta_cost_usd:.2f}$"
								)
							else:
								self.live_info_var.set(
								f"Previous block: {self.current_block_label} | +{delta_tokens} tokens | Cost: n/a (model price unknown)"
								)
						self.status_var.set(
							self._format_status_lines(
								"Block done",
								f"{processed}/{total}",
								f"Final errors: {failed}",
								paragraph_info,
							)
						)

					if self.debug_var.get() and debug_payload:
						source_bold_segments = int(debug_payload.get("source_bold_segments", 0) or 0)
						source_named_bold_runs = int(debug_payload.get("source_named_bold_runs", 0) or 0)
						source_has_named_bold = bool(debug_payload.get("source_has_named_bold", False))
						source_colored_segments = int(debug_payload.get("source_colored_segments", 0) or 0)
						output_bold_segments = int(debug_payload.get("output_bold_segments", 0) or 0)
						used_run_styles = debug_payload.get("used_run_styles", [])
						used_char_styles = debug_payload.get("used_char_styles", [])
						used_color_values = debug_payload.get("used_color_values", [])
						used_run_styles_text = ", ".join(str(s) for s in used_run_styles) if used_run_styles else "-"
						used_char_styles_text = ", ".join(str(s) for s in used_char_styles) if used_char_styles else "-"
						used_color_values_text = ", ".join(str(v) for v in used_color_values) if used_color_values else "-"
						self._append_error_log(
							"Debug Paragraph",
							[
								f"Paragraph: {debug_payload.get('paragraph', paragraph_info)}",
								f"Hard Tokens: {debug_payload.get('hard_token_count', '-')}",
								f"Color Tokens: {debug_payload.get('color_token_count', '-')}",
								f"Used Run Styles (name<style_id>): {used_run_styles_text}",
								f"Used Character Styles (name<style_id>): {used_char_styles_text}",
								f"Used Color Values (val/theme/tint/shade): {used_color_values_text}",
								f"Source Bold Runs: {debug_payload.get('source_bold_runs', '-')}",
								f"Source Named Bold Runs (Fett/Bold): {source_named_bold_runs}",
								f"Source Has Named Bold: {source_has_named_bold}",
								f"Source Bold Segments: {source_bold_segments}",
								f"Source Colored Segments: {source_colored_segments}",
								f"Preferred Bold Style: {debug_payload.get('preferred_style', '-')}",
								f"Style Fett Exists: {debug_payload.get('style_fett_exists', '-')}",
								f"Style Bold Exists: {debug_payload.get('style_bold_exists', '-')}",
								f"Output Bold Segments (**...**): {output_bold_segments}",
								f"Applied By Style: {debug_payload.get('applied_by_style', '-')}",
								f"Applied Direct Bold Fallback: {debug_payload.get('applied_direct', '-')}",
								f"Attempt: {attempt if attempt > 0 else '-'}",
							],
						)

				elif event == "paused_state":
					_, paused, processed, total, out_path = item
					self.is_paused = bool(paused)
					if self.is_running:
						self.pause_button.configure(text=("Resume" if self.is_paused else "Pause"), state="normal")
					if self.is_paused:
						self.status_var.set(self._format_status_lines("Paused", f"{processed}/{total}", "Output checkpoint saved"))
						checkpoint = Path(str(out_path))
						self.last_output = checkpoint
						self.last_output_exists = checkpoint.exists()
						if self.last_output_exists:
							self.reveal_button.grid()
					if self.is_paused:
						self._taskbar_update(state=ProgressType.PAUSED)
					else:
						self._taskbar_update(state=ProgressType.NORMAL)
				elif event == "done":
					_, total, translated, failed, warning_count, out_path, cancelled, failed_items, usage_summary = item
					self.warning_count_live = int(warning_count)
					self.error_count_live = int(failed)
					self._refresh_counters()
					if isinstance(usage_summary, dict):
						self.usage_prompt_tokens_live = int(usage_summary.get("prompt_tokens", 0) or 0)
						self.usage_completion_tokens_live = int(usage_summary.get("completion_tokens", 0) or 0)
						self.usage_total_tokens_live = int(usage_summary.get("total_tokens", 0) or 0)
						self.usage_pricing_known_live = bool(usage_summary.get("pricing_known", False))
						estimated_usd = usage_summary.get("estimated_cost_usd")
						if self.usage_pricing_known_live and isinstance(estimated_usd, (int, float)):
							self.usage_cost_usd_live = float(estimated_usd)
						self.usage_eur_rate_live = _coerce_float(usage_summary.get("eur_rate"))
						estimated_eur = usage_summary.get("estimated_cost_eur")
						if self.usage_pricing_known_live and isinstance(estimated_eur, (int, float)):
							self.usage_cost_eur_live = float(estimated_eur)
						self._refresh_usage_counters()
					self.progress_var.set(100)
					self.last_output = Path(out_path) if out_path else None
					self.last_output_exists = self.last_output is not None and self.last_output.exists()
					summary_col = f"Errors: {failed} | Warnings: {warning_count}"
					if cancelled:
						self.status_var.set(self._format_status_lines("Cancelled", f"{translated}/{total}", summary_col))
					elif total == 0 and translated == 0 and failed == 0:
						self.status_var.set(self._format_status_lines("Done", "No changes", summary_col))
						self._taskbar_flash_done()
					else:
						self.status_var.set(self._format_status_lines("Done", f"{translated}/{total}", summary_col))
						self._taskbar_flash_done()
					self.is_paused = False
					self._taskbar_update(state=ProgressType.NOPROGRESS)
					self._set_running(False)
					if self.last_output_exists:
						self.reveal_button.grid()
					else:
						self.reveal_button.grid_remove()
					for failed_item in failed_items:
						if str(failed_item.get("paragraph", "")) != "Snapshot":
							continue
						snapshot_msg = str(failed_item.get("error", "")).strip()
						if not snapshot_msg:
							continue
						self._append_warning_log("Snapshot not saved", [snapshot_msg])

				elif event == "models_loaded":
					_, model_ids = item
					self._set_model_choices(model_ids)
					self._save_settings()
					self.status_var.set(f"Models loaded ({len(self.model_choices)}). Ready.")
					self._set_models_loading(False)
					if not self.is_running and not self.prices_loading:
						self._taskbar_update(state=ProgressType.NOPROGRESS)

				elif event == "models_error":
					_, msg, details, on_startup = item
					self._set_models_loading(False)
					self._append_error_log(
						"Error loading models",
						[
							f"Message: {msg}",
							"Hint: Check the API key or network connection.",
							"Details: See stack trace in the developer terminal.",
						],
					)
					prefix = "At startup" if on_startup else "During model refresh"
					self.status_var.set(f"{prefix} failed. Using fallback models.")
					if not self.is_running and not self.prices_loading:
						self._taskbar_update(state=ProgressType.ERROR)
				elif event == "prices_progress":
					_, done, total, message = item
					try:
						done_steps = max(0, int(done))
					except Exception:
						done_steps = 0
					try:
						total_steps = max(0, int(total))
					except Exception:
						total_steps = 0
					if total_steps > 0:
						percent = min(100.0, (done_steps / float(total_steps)) * 100.0)
					else:
						percent = 0.0
					if percent <= 0.0 and done_steps < total_steps:
						percent = 1.0
					if self.prices_loading and not self.is_running:
						self.progress_var.set(percent)
						self._taskbar_update(state=ProgressType.NORMAL, percent=percent)
					if message:
						self.status_var.set(str(message))
				elif event == "prices_updated":
					_, summary = item
					self.progress_var.set(100)
					refresh_model_pricing_cache()
					model_ids_from_update = list(summary.get("model_ids", [])) if isinstance(summary, dict) else []
					if model_ids_from_update:
						self._set_model_choices(model_ids_from_update)
					else:
						self._set_model_choices(self.model_choices)
					self._set_prices_loading(False)
					resolved_count = int(summary.get("resolved_count", 0) or 0) if isinstance(summary, dict) else 0
					total_count = int(summary.get("total_count", 0) or 0) if isinstance(summary, dict) else 0
					updated_count = int(summary.get("updated_count", 0) or 0) if isinstance(summary, dict) else 0
					seeded_count = int(summary.get("seeded_count", 0) or 0) if isinstance(summary, dict) else 0
					unresolved_models = list(summary.get("unresolved_models", [])) if isinstance(summary, dict) else []
					diagnostics = list(summary.get("diagnostics", [])) if isinstance(summary, dict) else []

					info_lines = [
						f"Updated entries: {updated_count}",
						f"Seeded entries: {seeded_count}",
						f"Resolved prices: {resolved_count}/{total_count}",
					]
					if unresolved_models:
						preview = ", ".join(unresolved_models[:12])
						if len(unresolved_models) > 12:
							preview = f"{preview}, ..."
						info_lines.append(f"Unresolved models: {preview}")
					for diag in diagnostics:
						info_lines.append(f"Detail: {diag}")
					self._append_info_log("Model prices updated", info_lines)
					self.status_var.set(f"Model prices updated ({resolved_count}/{total_count} resolved).")
					self._taskbar_flash_done()
					if not self.is_running and not self.models_loading:
						self._taskbar_update(state=ProgressType.NOPROGRESS)
				elif event == "prices_error":
					_, msg, details = item
					self._set_prices_loading(False)
					self.progress_var.set(0)
					self._append_error_log(
						"Error updating model prices",
						[
							f"Message: {msg}",
							f"Traceback: {details.splitlines()[-1] if details else '-'}",
						],
					)
					self.status_var.set("Model price update failed.")
					if not self.is_running and not self.models_loading:
						self._taskbar_update(state=ProgressType.ERROR)
				elif event == "styles_scanned":
					_, styles, char_styles = item
					self.available_styles = styles
					self.available_char_styles = char_styles
					self.styles_scanning = False
					existing = self.selected_styles & set(styles)
					if existing:
						self.selected_styles = existing
					else:
						self.selected_styles = {s for s in styles if s in TARGET_STYLE_NAMES}
						if not self.selected_styles and styles:
							self.selected_styles = {styles[0]}
					existing_char = self.selected_char_styles & set(char_styles)
					if existing_char:
						self.selected_char_styles = existing_char
					elif "selected_char_styles" not in self.settings:
						# First run fallback: preselect detected non-default character styles.
						self.selected_char_styles = set(char_styles)
					else:
						self.selected_char_styles = set()
					self._update_styles_info()
					self._update_char_styles_info()
					self._sync_controls()
					self.status_var.set(f"{len(styles)} styles, {len(char_styles)} character styles found. Ready.")

				elif event == "styles_scan_error":
					_, msg, input_path_text, on_startup, is_locked = item
					self.styles_scanning = False
					if is_locked:
						self.styles_info_var.set("File open in Word")
						self.char_styles_info_var.set("File open in Word")
						self.status_var.set("The selected DOCX is open in Word. Please close it and reload.")
						self._append_error_log(
							"File locked",
							[
								f"File: {input_path_text}",
								"The file is open in Word (lock file detected).",
								"Please close Word and re-select the file.",
							],
						)
					elif on_startup:
						self.file_var.set("")
						self.styles_info_var.set("Select a file first")
						self.char_styles_info_var.set("Select a file first")
						self.status_var.set("Previously saved file is not available. Please re-select a file.")
						self._save_settings(input_file="")
					else:
						self.styles_info_var.set("Error reading styles")
						self.char_styles_info_var.set("Error reading character styles")
						self._append_error_log(
							"Error reading styles",
							[
								f"File: {input_path_text}",
								f"Message: {msg}",
							],
						)
					self._sync_controls()
				elif event == "error":
					_, exc, details = item
					self.error_count_live += 1
					self._refresh_counters()
					summary, detail_lines = format_openai_exception(exc)
					self.status_var.set("Translation error.")
					self.is_paused = False
					self._taskbar_update(state=ProgressType.ERROR)
					self._set_running(False)
					self._append_error_log("Runtime error", [*detail_lines, f"Traceback: {details.splitlines()[-1] if details else '-'}"])
		except Empty:
			pass
		finally:
			self.root.after(100, self._poll_queue)

def main() -> None:
	root = tk.Tk()
	root.tk.call("encoding", "system", "utf-8")
	TranslatorApp(root)
	root.mainloop()

if __name__ == "__main__":
	main()